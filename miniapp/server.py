# -*- coding: utf-8 -*-
"""
Telegram Mini App backend for the local LLM Telegram bot.

Отдельный локальный веб-сервер (FastAPI + uvicorn), который:
  * читает ту же SQLite-базу RAG-памяти, что и основной бот (rag_files/rag_memory.sqlite3);
  * общается с LM Studio по http://localhost:1234 (чат со стримингом, список/загрузка моделей);
  * принимает загрузку документов и добавляет их в общую RAG-память;
  * отдаёт REST API и одностраничный премиум-фронтенд (static/index.html) как Telegram Mini App;
  * проверяет подпись Telegram initData, чтобы понять, кто открыл приложение, и кто админ.

ВАЖНО: основной файл бота (telegram-local-llm-bot.py) НЕ импортируется и НЕ изменяется.

Запуск:
    python miniapp/server.py
или (из корня проекта):
    python -m uvicorn miniapp.server:app --host 0.0.0.0 --port 8080
"""

import hashlib
import hmac
import json
import os
import re
import sqlite3
import tempfile
import threading
import time
from typing import Optional
from urllib.parse import parse_qsl

import numpy as np
import requests
from dotenv import load_dotenv
from fastapi import Body, FastAPI, File, Header, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

# ------------------------------------------------------------------
# Пути и конфигурация.
# ------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
load_dotenv(os.path.join(BASE_DIR, ".env"))


def project_path(path: str) -> str:
    if os.path.isabs(path):
        return path
    return os.path.join(BASE_DIR, path)


RAG_STORAGE_DIR = project_path(os.environ.get("RAG_STORAGE_DIR", "rag_files"))
RAG_DB_PATH = project_path(
    os.environ.get("RAG_DB_PATH", os.path.join(RAG_STORAGE_DIR, "rag_memory.sqlite3"))
)
RAG_FILES_DIR = RAG_STORAGE_DIR

RAG_EMBEDDING_MODEL_NAME = os.environ.get(
    "RAG_EMBEDDING_MODEL",
    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
)
RAG_TOP_K = int(os.environ.get("RAG_TOP_K", "4"))
RAG_MAX_CONTEXT_CHARS = int(os.environ.get("RAG_MAX_CONTEXT_CHARS", "4500"))
RAG_MIN_SCORE = float(os.environ.get("RAG_MIN_SCORE", "0.22"))
RAG_KEYWORD_WEIGHT = float(os.environ.get("RAG_KEYWORD_WEIGHT", "0.45"))
RAG_SEMANTIC_WEIGHT = float(os.environ.get("RAG_SEMANTIC_WEIGHT", "0.55"))
RAG_CHUNK_SIZE = int(os.environ.get("RAG_CHUNK_SIZE", "900"))
RAG_CHUNK_OVERLAP = int(os.environ.get("RAG_CHUNK_OVERLAP", "150"))
RAG_MAX_FILE_MB = int(os.environ.get("RAG_MAX_FILE_MB", "20"))
SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".md", ".log", ".csv", ".json", ".docx"}

LM_STUDIO_BASE_URL = os.environ.get("LM_STUDIO_BASE_URL", "http://localhost:1234")
LM_STUDIO_CHAT_URL = f"{LM_STUDIO_BASE_URL}/v1/chat/completions"
LM_STUDIO_MODELS_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models"
LM_STUDIO_LOAD_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models/load"
LM_STUDIO_UNLOAD_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models/unload"

SYSTEM_PROMPT = os.environ.get(
    "SYSTEM_PROMPT",
    "Ты — полезный ассистент. Отвечай ТОЛЬКО на русском языке. Давай краткие, естественные ответы. "
    "Используй эмодзи и Markdown-форматирование (жирный, списки, заголовки, код) там, где это уместно, "
    "как современные чат-ассистенты.",
)

DEFAULT_MODEL_KEY = "balanced"
AVAILABLE_MODELS = {
    "fast": {"id": "google/gemma-4-e4b", "label": "⚡ Быстрая (117 т/с)"},
    "balanced": {"id": "google/gemma-4-12b-qat", "label": "⚖️ Сбалансированная (70 т/с)"},
    "smart": {"id": "google/gemma-4-26b-a4b-qat", "label": "🧠 Умная (33 т/с)"},
}

TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
ADMIN_TELEGRAM_IDS = {
    item.strip()
    for item in os.environ.get("ADMIN_TELEGRAM_IDS", "").split(",")
    if item.strip()
}
REQUIRE_TELEGRAM_AUTH = os.environ.get(
    "MINIAPP_REQUIRE_AUTH", "false"
).strip().lower() in {"1", "true", "yes", "on"}
DEV_CHAT_ID = os.environ.get("MINIAPP_DEV_CHAT_ID", "dev-local")

conversation_history: dict = {}
webapp_selected_model: dict = {}
# Статус фоновой загрузки модели по chat_id: {"key":..., "state":"loading|ready|error", "error":...}
model_load_status: dict = {}

app = FastAPI(title="Telegram Bot Mini App")

# ------------------------------------------------------------------
# Ленивая загрузка модели эмбеддингов.
# ------------------------------------------------------------------
_embedding_model = None


def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        from sentence_transformers import SentenceTransformer

        _embedding_model = SentenceTransformer(RAG_EMBEDDING_MODEL_NAME)
    return _embedding_model


def embed_texts(texts):
    model = get_embedding_model()
    vectors = model.encode(texts, normalize_embeddings=True, convert_to_numpy=True)
    return np.asarray(vectors, dtype=np.float32)


# Ленивая загрузка Whisper (для голосового ввода).
_whisper_model = None


def get_whisper():
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel

        name = os.environ.get("WHISPER_MODEL", "medium")
        _whisper_model = WhisperModel(name, device="cpu", compute_type="int8")
    return _whisper_model


# ------------------------------------------------------------------
# База данных (та же схема, что и у бота).
# ------------------------------------------------------------------
def db_connect():
    if not os.path.exists(RAG_DB_PATH):
        os.makedirs(os.path.dirname(os.path.abspath(RAG_DB_PATH)), exist_ok=True)
    return sqlite3.connect(RAG_DB_PATH)


_schema_ready = False


def ensure_min_schema():
    """CREATE TABLE/ALTER здесь идемпотентны, поэтому реально прогоняем их один раз за процесс."""
    global _schema_ready
    if _schema_ready:
        return
    os.makedirs(os.path.dirname(os.path.abspath(RAG_DB_PATH)), exist_ok=True)
    with db_connect() as conn:
        conn.execute(
            """CREATE TABLE IF NOT EXISTS rag_settings (
                chat_id TEXT PRIMARY KEY,
                rag_only INTEGER NOT NULL DEFAULT 0,
                updated_at INTEGER NOT NULL
            )"""
        )
        cols = {r[1] for r in conn.execute("PRAGMA table_info(rag_settings)").fetchall()}
        if "study_mode" not in cols:
            conn.execute("ALTER TABLE rag_settings ADD COLUMN study_mode INTEGER NOT NULL DEFAULT 0")
        conn.execute(
            """CREATE TABLE IF NOT EXISTS webapp_chat_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at INTEGER NOT NULL
            )"""
        )
        mcols = {r[1] for r in conn.execute("PRAGMA table_info(webapp_chat_messages)").fetchall()}
        if "session_id" not in mcols:
            conn.execute("ALTER TABLE webapp_chat_messages ADD COLUMN session_id INTEGER")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_webapp_chat ON webapp_chat_messages(chat_id, id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_webapp_msg_sess ON webapp_chat_messages(session_id, id)")
        conn.execute(
            """CREATE TABLE IF NOT EXISTS webapp_chat_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT NOT NULL,
                title TEXT,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL
            )"""
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_webapp_sessions ON webapp_chat_sessions(chat_id, updated_at)")
        conn.commit()
    _schema_ready = True


# ---- сессии (беседы) ----
def create_session(chat_id, title=None):
    ensure_min_schema()
    now = int(time.time())
    with db_connect() as conn:
        cur = conn.execute(
            "INSERT INTO webapp_chat_sessions (chat_id, title, created_at, updated_at) VALUES (?, ?, ?, ?)",
            (str(chat_id), title, now, now),
        )
        conn.commit()
        return cur.lastrowid


def get_session(chat_id, session_id):
    with db_connect() as conn:
        r = conn.execute(
            "SELECT id, title FROM webapp_chat_sessions WHERE chat_id = ? AND id = ?",
            (str(chat_id), int(session_id)),
        ).fetchone()
    return {"id": r[0], "title": r[1]} if r else None


def touch_session(session_id):
    with db_connect() as conn:
        conn.execute("UPDATE webapp_chat_sessions SET updated_at = ? WHERE id = ?",
                     (int(time.time()), int(session_id)))
        conn.commit()


def set_session_title(session_id, title):
    with db_connect() as conn:
        conn.execute(
            "UPDATE webapp_chat_sessions SET title = ? WHERE id = ? AND (title IS NULL OR title = '')",
            (title[:80], int(session_id)),
        )
        conn.commit()


def list_sessions(chat_id):
    ensure_min_schema()
    with db_connect() as conn:
        rows = conn.execute(
            """SELECT s.id, s.title, s.created_at, s.updated_at,
                      (SELECT COUNT(*) FROM webapp_chat_messages m WHERE m.session_id = s.id) AS cnt
               FROM webapp_chat_sessions s WHERE s.chat_id = ?
               ORDER BY s.updated_at DESC""",
            (str(chat_id),),
        ).fetchall()
    return [
        {"id": r[0], "title": r[1] or "Новый чат", "created_at": r[2], "updated_at": r[3], "count": r[4]}
        for r in rows
    ]


def delete_session(chat_id, session_id):
    with db_connect() as conn:
        conn.execute("DELETE FROM webapp_chat_messages WHERE chat_id = ? AND session_id = ?",
                     (str(chat_id), int(session_id)))
        conn.execute("DELETE FROM webapp_chat_sessions WHERE chat_id = ? AND id = ?",
                     (str(chat_id), int(session_id)))
        conn.commit()


def save_chat_message(chat_id, session_id, role, content):
    ensure_min_schema()
    with db_connect() as conn:
        conn.execute(
            "INSERT INTO webapp_chat_messages (chat_id, session_id, role, content, created_at) "
            "VALUES (?, ?, ?, ?, ?)",
            (str(chat_id), int(session_id), role, content, int(time.time())),
        )
        conn.commit()


def load_session_messages(chat_id, session_id, limit=300):
    ensure_min_schema()
    with db_connect() as conn:
        rows = conn.execute(
            "SELECT role, content, created_at FROM webapp_chat_messages "
            "WHERE chat_id = ? AND session_id = ? ORDER BY id ASC LIMIT ?",
            (str(chat_id), int(session_id), int(limit)),
        ).fetchall()
    return [{"role": r[0], "content": r[1], "created_at": r[2]} for r in rows]


def delete_last_messages(session_id, n):
    ensure_min_schema()
    with db_connect() as conn:
        conn.execute(
            "DELETE FROM webapp_chat_messages WHERE id IN "
            "(SELECT id FROM webapp_chat_messages WHERE session_id = ? ORDER BY id DESC LIMIT ?)",
            (int(session_id), int(n)),
        )
        conn.commit()


def document_row_to_dict(row):
    if not row:
        return None
    return {
        "id": row[0],
        "source_name": row[1],
        "file_size": row[2] or 0,
        "file_ext": row[3] or "",
        "chunks_count": row[4] or 0,
        "summary": row[5] or "",
        "created_at": row[6] or 0,
        "updated_at": row[7] or 0,
        "saved_path": row[8] or "",
        "source_hash": row[9] or "",
    }


_DOC_COLS = """id, source_name, file_size, file_ext, chunks_count, summary,
               created_at, updated_at, saved_path, source_hash"""


def get_rag_documents(chat_id):
    with db_connect() as conn:
        rows = conn.execute(
            f"SELECT {_DOC_COLS} FROM rag_documents WHERE chat_id = ? "
            "ORDER BY updated_at DESC, id DESC",
            (str(chat_id),),
        ).fetchall()
    return [document_row_to_dict(r) for r in rows]


def get_document_by_id(chat_id, doc_id):
    with db_connect() as conn:
        row = conn.execute(
            f"SELECT {_DOC_COLS} FROM rag_documents WHERE chat_id = ? AND id = ?",
            (str(chat_id), int(doc_id)),
        ).fetchone()
    return document_row_to_dict(row)


def get_document_by_hash(chat_id, file_hash):
    with db_connect() as conn:
        row = conn.execute(
            f"SELECT {_DOC_COLS} FROM rag_documents WHERE chat_id = ? AND source_hash = ?",
            (str(chat_id), file_hash),
        ).fetchone()
    return document_row_to_dict(row)


def get_document_chunks(chat_id, doc_id, limit=None):
    sql = ("SELECT chunk_index, text FROM rag_chunks "
           "WHERE chat_id = ? AND document_id = ? ORDER BY chunk_index ASC")
    params = [str(chat_id), int(doc_id)]
    if limit:
        sql += " LIMIT ?"
        params.append(int(limit))
    with db_connect() as conn:
        return conn.execute(sql, params).fetchall()


def delete_rag_document_by_id(chat_id, doc_id):
    doc = get_document_by_id(chat_id, doc_id)
    if not doc:
        return False
    with db_connect() as conn:
        conn.execute("DELETE FROM rag_chunks WHERE chat_id = ? AND document_id = ?",
                     (str(chat_id), int(doc_id)))
        conn.execute("DELETE FROM rag_chunks WHERE chat_id = ? AND source_name = ?",
                     (str(chat_id), doc["source_name"]))
        conn.execute("DELETE FROM rag_documents WHERE chat_id = ? AND id = ?",
                     (str(chat_id), int(doc_id)))
        conn.commit()
    saved_path = doc.get("saved_path")
    if saved_path and os.path.exists(saved_path):
        try:
            os.remove(saved_path)
        except OSError:
            pass
    return True


def get_rag_stats(chat_id):
    with db_connect() as conn:
        total_chunks = conn.execute(
            "SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ?", (str(chat_id),)
        ).fetchone()[0]
    return total_chunks, get_rag_documents(chat_id)


def get_settings(chat_id):
    ensure_min_schema()
    with db_connect() as conn:
        row = conn.execute(
            "SELECT rag_only, study_mode FROM rag_settings WHERE chat_id = ?",
            (str(chat_id),),
        ).fetchone()
    if not row:
        return {"rag_only": False, "study_mode": False}
    return {"rag_only": bool(row[0]), "study_mode": bool(row[1])}


def set_setting(chat_id, key, value):
    if key not in {"rag_only", "study_mode"}:
        raise ValueError("unknown setting")
    ensure_min_schema()
    now = int(time.time())
    with db_connect() as conn:
        conn.execute(
            "INSERT INTO rag_settings (chat_id, rag_only, study_mode, updated_at) "
            "VALUES (?, ?, ?, ?) "
            "ON CONFLICT(chat_id) DO UPDATE SET {c}=excluded.{c}, updated_at=excluded.updated_at".format(c=key),
            (
                str(chat_id),
                1 if (key == "rag_only" and value) else 0,
                1 if (key == "study_mode" and value) else 0,
                now,
            ),
        )
        conn.commit()
    return get_settings(chat_id)


def get_db_size_text():
    try:
        size = os.path.getsize(RAG_DB_PATH)
    except OSError:
        return "0 B"
    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def get_admin_stats():
    with db_connect() as conn:
        users_count = conn.execute("SELECT COUNT(*) FROM bot_users").fetchone()[0]
        chats_with_chunks = conn.execute("SELECT COUNT(DISTINCT chat_id) FROM rag_chunks").fetchone()[0]
        total_chunks = conn.execute("SELECT COUNT(*) FROM rag_chunks").fetchone()[0]
        total_sources = conn.execute(
            "SELECT COUNT(*) FROM (SELECT chat_id, source_name FROM rag_chunks "
            "GROUP BY chat_id, source_name)"
        ).fetchone()[0]
        rag_only_count = conn.execute("SELECT COUNT(*) FROM rag_settings WHERE rag_only = 1").fetchone()[0]
    return {
        "users_count": users_count,
        "chats_with_chunks": chats_with_chunks,
        "total_chunks": total_chunks,
        "total_sources": total_sources,
        "rag_only_count": rag_only_count,
        "db_size": get_db_size_text(),
        "db_path": RAG_DB_PATH,
    }


def get_admin_users(limit=30):
    with db_connect() as conn:
        rows = conn.execute(
            "SELECT chat_id, user_id, username, first_name, last_name, last_seen_at, message_count "
            "FROM bot_users ORDER BY last_seen_at DESC LIMIT ?",
            (int(limit),),
        ).fetchall()
    return [
        {
            "chat_id": r[0], "user_id": r[1], "username": r[2], "first_name": r[3],
            "last_name": r[4], "last_seen_at": r[5], "message_count": r[6],
        }
        for r in rows
    ]


# ------------------------------------------------------------------
# RAG-поиск (гибрид: semantic + keyword).
# ------------------------------------------------------------------
STOPWORDS = {
    "и", "в", "во", "не", "что", "он", "на", "я", "с", "со", "как", "а", "то",
    "все", "она", "так", "его", "но", "да", "ты", "к", "у", "же", "вы", "за",
    "бы", "по", "только", "ее", "мне", "было", "вот", "от", "меня", "о", "из",
    "the", "a", "an", "of", "to", "in", "is", "it", "and", "or",
}


def extract_keywords(query):
    tokens = re.findall(r"[\wа-яА-ЯёЁ\-]+", (query or "").lower())
    return [t for t in tokens if len(t) >= 3 and t not in STOPWORDS]


def calculate_keyword_score(query, text):
    query = (query or "").strip().lower()
    text_lower = (text or "").lower()
    if not query or not text_lower:
        return 0.0
    keywords = extract_keywords(query)
    if not keywords:
        return 0.0
    matches = important = 0
    for token in keywords:
        if token in text_lower:
            matches += 1
            if any(ch.isdigit() for ch in token) or "-" in token or "_" in token:
                important += 1
    token_score = matches / max(1, len(keywords))
    important_boost = min(0.25, important * 0.08)
    phrase_boost = 0.35 if (len(query) >= 10 and query in text_lower) else 0.0
    return min(1.0, token_score * 0.75 + important_boost + phrase_boost)


def rank_rag_chunks(chat_id, query, top_k=RAG_TOP_K):
    if not isinstance(query, str) or len(query.strip()) < 3:
        return []
    with db_connect() as conn:
        rows = conn.execute(
            "SELECT source_name, chunk_index, text, embedding FROM rag_chunks WHERE chat_id = ?",
            (str(chat_id),),
        ).fetchall()
    if not rows:
        return []
    query_embedding = embed_texts([query])[0]
    ranked = []
    for source_name, chunk_index, text, emb_blob in rows:
        emb = np.frombuffer(emb_blob, dtype=np.float32)
        semantic = float(np.dot(query_embedding, emb)) if emb.size == query_embedding.size else 0.0
        keyword = calculate_keyword_score(query, text)
        hybrid = semantic * RAG_SEMANTIC_WEIGHT + keyword * RAG_KEYWORD_WEIGHT
        if keyword >= 0.65:
            hybrid = max(hybrid, keyword)
        if semantic >= RAG_MIN_SCORE or keyword >= 0.35 or hybrid >= RAG_MIN_SCORE:
            ranked.append({
                "score": hybrid, "semantic_score": semantic, "keyword_score": keyword,
                "source_name": source_name, "chunk_index": chunk_index, "text": text,
            })
    ranked.sort(key=lambda x: x["score"], reverse=True)
    return ranked[:top_k]


def search_rag_with_sources(chat_id, query, top_k=RAG_TOP_K):
    selected = rank_rag_chunks(chat_id, query, top_k=top_k)
    if not selected:
        return "", []
    parts, used, total = [], [], 0
    for item in selected:
        block = (f"[Источник: {item['source_name']}, фрагмент {item['chunk_index'] + 1}, "
                 f"релевантность {item['score']:.2f}]\n{item['text']}")
        if total + len(block) > RAG_MAX_CONTEXT_CHARS:
            remaining = RAG_MAX_CONTEXT_CHARS - total
            if remaining <= 300:
                break
            block = block[:remaining].rstrip() + "..."
        parts.append(block)
        used.append(item)
        total += len(block)
    return "\n\n---\n\n".join(parts), used


def build_rag_system_message(rag_context, rag_only=False):
    if rag_only:
        return ("Ты работаешь в режиме RAG-only. Отвечай ТОЛЬКО на основании фрагментов документов ниже.\n"
                "Если в этих фрагментах нет ответа, прямо скажи: 'В загруженных документах я не нашёл ответа.'\n"
                "Не используй внешние знания и не додумывай факты.\n\n"
                f"КОНТЕКСТ ИЗ ДОКУМЕНТОВ:\n{rag_context}")
    return ("Ниже есть релевантные фрагменты из документов, которые пользователь ранее загрузил в память.\n"
            "Используй их только если они действительно помогают ответить на текущий вопрос.\n"
            "Если в документах нет ответа или информации мало, честно скажи об этом.\n"
            "Не выдумывай факты сверх предоставленного контекста.\n\n"
            f"КОНТЕКСТ ИЗ ДОКУМЕНТОВ:\n{rag_context}")


STUDY_MODE_PROMPT = (
    "Отвечай в режиме обучения: объясняй пошагово, приводи примеры и в конце задавай "
    "один короткий вопрос для проверки понимания."
)


# ------------------------------------------------------------------
# Загрузка/извлечение документов (совместимо с ботом).
# ------------------------------------------------------------------
def clean_text(text):
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text_into_chunks(text, chunk_size=RAG_CHUNK_SIZE, overlap=RAG_CHUNK_OVERLAP):
    text = clean_text(text)
    if not text:
        return []
    compact = re.sub(r"\s+", " ", text).strip()
    chunks, start = [], 0
    while start < len(compact):
        end = min(len(compact), start + chunk_size)
        if end < len(compact):
            sentence_end = max(
                compact.rfind(". ", start, end),
                compact.rfind("! ", start, end),
                compact.rfind("? ", start, end),
                compact.rfind("; ", start, end),
            )
            if sentence_end > start + int(chunk_size * 0.55):
                end = sentence_end + 1
        chunk = compact[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(compact):
            break
        start = max(0, end - overlap)
        if start >= end:
            start = end
    return chunks


def _extract_pdf(path):
    errors = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(path, strict=False)
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as e:
                errors.append(str(e))
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception as e:
        errors.append(f"pypdf: {e}")
    try:
        import fitz

        doc = fitz.open(path)
        parts = []
        for i in range(len(doc)):
            try:
                parts.append(doc.load_page(i).get_text("text") or "")
            except Exception as e:
                errors.append(str(e))
        doc.close()
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception as e:
        errors.append(f"fitz: {e}")
    raise RuntimeError("Не удалось извлечь текст из PDF (скан/битый/без текста). " + " | ".join(errors[:2]))


def _extract_docx(path):
    from docx import Document

    document = Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_plain(path):
    with open(path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_text_from_file(path, file_name):
    ext = os.path.splitext(file_name.lower())[1]
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in {".txt", ".md", ".log", ".csv", ".json"}:
        return _extract_plain(path)
    raise ValueError("Неподдерживаемый формат. Разрешены: PDF, TXT, MD, LOG, CSV, JSON, DOCX.")


def sanitize_filename(name):
    name = os.path.basename(name or "document")
    name = re.sub(r"[^a-zA-Zа-яА-ЯёЁ0-9._()\- ]+", "_", name).strip()
    return name or "document"


def save_original_document(chat_id, file_name, file_bytes, file_hash):
    safe = sanitize_filename(file_name)
    chat_dir = os.path.join(RAG_FILES_DIR, f"chat_{chat_id}")
    os.makedirs(chat_dir, exist_ok=True)
    base, ext = os.path.splitext(safe)
    saved_path = os.path.join(chat_dir, f"{int(time.time())}_{file_hash[:10]}_{base[:70]}{ext}")
    with open(saved_path, "wb") as f:
        f.write(file_bytes)
    return saved_path


def add_document_to_rag(chat_id, source_name, text, source_hash, saved_path, file_size, file_ext):
    text = clean_text(text)
    chunks = split_text_into_chunks(text)
    if not chunks:
        return None, 0, 0
    source_name = sanitize_filename(source_name)
    now = int(time.time())
    embeddings = embed_texts(chunks)
    inserted = 0
    with db_connect() as conn:
        conn.execute(
            "INSERT INTO rag_documents "
            "(chat_id, source_name, source_hash, saved_path, file_size, file_ext, chunks_count, summary, created_at, updated_at) "
            "VALUES (?, ?, ?, ?, ?, ?, 0, NULL, ?, ?) "
            "ON CONFLICT(chat_id, source_hash) DO UPDATE SET "
            "source_name=excluded.source_name, "
            "saved_path=COALESCE(excluded.saved_path, rag_documents.saved_path), "
            "file_size=CASE WHEN excluded.file_size>0 THEN excluded.file_size ELSE rag_documents.file_size END, "
            "file_ext=excluded.file_ext, updated_at=excluded.updated_at",
            (str(chat_id), source_name, source_hash, saved_path, int(file_size or 0), file_ext or "", now, now),
        )
        doc_id = conn.execute(
            "SELECT id FROM rag_documents WHERE chat_id = ? AND source_hash = ?",
            (str(chat_id), source_hash),
        ).fetchone()[0]
        for index, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            chunk_hash = hashlib.sha256(f"{source_hash}:{index}:{chunk}".encode("utf-8")).hexdigest()
            try:
                conn.execute(
                    "INSERT INTO rag_chunks "
                    "(chat_id, document_id, source_name, chunk_index, chunk_hash, text, embedding, created_at) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (str(chat_id), doc_id, source_name, index, chunk_hash, chunk,
                     emb.astype(np.float32).tobytes(), now),
                )
                inserted += 1
            except sqlite3.IntegrityError:
                pass
        cnt = conn.execute(
            "SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ? AND document_id = ?",
            (str(chat_id), doc_id),
        ).fetchone()[0]
        conn.execute("UPDATE rag_documents SET chunks_count = ?, updated_at = ? WHERE id = ?",
                     (cnt, now, doc_id))
        conn.commit()
    return doc_id, inserted, len(chunks)


# ------------------------------------------------------------------
# LM Studio.
# ------------------------------------------------------------------
def get_headers():
    headers = {"Content-Type": "application/json"}
    if os.environ.get("LM_STUDIO_API_KEY"):
        headers["Authorization"] = f"Bearer {os.environ.get('LM_STUDIO_API_KEY')}"
    return headers


def get_loaded_models():
    try:
        resp = requests.get(LM_STUDIO_MODELS_URL, headers=get_headers(), timeout=10)
        if resp.status_code != 200:
            return []
        data = resp.json()
        loaded = []
        for item in data.get("models", []):
            for instance in item.get("loaded_instances", []):
                loaded.append(instance.get("id", item.get("key")))
        return loaded
    except Exception:
        return []


def lm_reachable():
    try:
        requests.get(LM_STUDIO_MODELS_URL, headers=get_headers(), timeout=3)
        return True
    except Exception:
        return False


def unload_model_instance(instance_id):
    try:
        requests.post(LM_STUDIO_UNLOAD_URL, headers=get_headers(),
                      json={"instance_id": instance_id}, timeout=8)
    except Exception:
        pass


def ensure_model_loaded(model_id):
    """
    Загружает выбранную модель СРАЗУ (не дожидаясь сообщения).
    Способ 1: явный endpoint загрузки LM Studio (если поддерживается версией).
    Способ 2 (надёжный fallback): крошечный chat-запрос — заставляет LM Studio
    выполнить JIT-загрузку модели, ровно как при обычном сообщении.
    Плюс пытаемся выгрузить остальные известные модели.
    """
    if not lm_reachable():
        return False, "LM Studio недоступен"

    loaded = get_loaded_models()
    known = {info["id"] for info in AVAILABLE_MODELS.values()}

    # Выгружаем ТОЛЬКО реально загруженные прочие модели (без ошибок model_not_found).
    for m in loaded:
        if m in known and m != model_id:
            unload_model_instance(m)

    # Если целевая модель уже в памяти — переключение мгновенное.
    if model_id in loaded:
        return True, None

    # Иначе форсируем JIT-загрузку минимальным chat-запросом.
    try:
        r = requests.post(
            LM_STUDIO_CHAT_URL, headers=get_headers(),
            json={"model": model_id,
                  "messages": [{"role": "user", "content": "."}],
                  "max_tokens": 1, "stream": False},
            timeout=180,
        )
        if r.status_code == 200:
            return True, None
        return False, f"Не удалось загрузить (HTTP {r.status_code})"
    except requests.exceptions.ConnectionError:
        return False, "LM Studio недоступен"
    except Exception as e:
        return False, str(e)


# ------------------------------------------------------------------
# Аутентификация Telegram Mini App (валидация initData по HMAC).
# ------------------------------------------------------------------
def validate_init_data(init_data: str):
    if not init_data or not TELEGRAM_BOT_TOKEN:
        return None
    try:
        pairs = dict(parse_qsl(init_data, keep_blank_values=True))
    except Exception:
        return None
    received_hash = pairs.pop("hash", None)
    if not received_hash:
        return None
    data_check_string = "\n".join(f"{k}={pairs[k]}" for k in sorted(pairs.keys()))
    secret_key = hmac.new(b"WebAppData", TELEGRAM_BOT_TOKEN.encode(), hashlib.sha256).digest()
    calc_hash = hmac.new(secret_key, data_check_string.encode(), hashlib.sha256).hexdigest()
    if not hmac.compare_digest(calc_hash, received_hash):
        return None
    user = None
    if "user" in pairs:
        try:
            user = json.loads(pairs["user"])
        except Exception:
            user = None
    return {"user": user, "auth_date": pairs.get("auth_date")}


def resolve_identity(init_data: Optional[str]):
    parsed = validate_init_data(init_data) if init_data else None
    if parsed and parsed.get("user"):
        uid = str(parsed["user"].get("id"))
        return uid, uid in ADMIN_TELEGRAM_IDS, parsed["user"]
    if REQUIRE_TELEGRAM_AUTH:
        raise HTTPException(status_code=401, detail="Требуется открытие через Telegram")
    return DEV_CHAT_ID, True, {"id": DEV_CHAT_ID, "first_name": "Dev", "username": "dev"}


def require_admin(is_admin: bool):
    if not is_admin:
        raise HTTPException(status_code=403, detail="Только для администраторов")


def telegram_bot_configured():
    if not TELEGRAM_BOT_TOKEN:
        return None
    try:
        r = requests.get(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getMe", timeout=5)
        return r.status_code == 200 and r.json().get("ok", False)
    except Exception:
        return False


# ------------------------------------------------------------------
# API.
# ------------------------------------------------------------------
@app.get("/api/health")
def api_health():
    return {"ok": True, "lm_studio_reachable": lm_reachable(), "db_exists": os.path.exists(RAG_DB_PATH)}


@app.get("/api/me")
def api_me(x_init_data: Optional[str] = Header(default=None)):
    chat_id, is_admin, user = resolve_identity(x_init_data)
    return {
        "chat_id": chat_id,
        "is_admin": is_admin,
        "user": user,
        "selected_model": webapp_selected_model.get(chat_id, DEFAULT_MODEL_KEY),
        "settings": get_settings(chat_id),
    }


@app.get("/api/overview")
def api_overview(x_init_data: Optional[str] = Header(default=None)):
    chat_id, is_admin, user = resolve_identity(x_init_data)
    total_chunks, docs = get_rag_stats(chat_id)
    settings = get_settings(chat_id)
    loaded = get_loaded_models()
    model_key = webapp_selected_model.get(chat_id, DEFAULT_MODEL_KEY)
    active = AVAILABLE_MODELS.get(model_key, {})
    return {
        "user": user,
        "is_admin": is_admin,
        "documents": len(docs),
        "chunks": total_chunks,
        "rag_only": settings["rag_only"],
        "study_mode": settings["study_mode"],
        "active_model": {"key": model_key, "label": active.get("label", ""), "id": active.get("id", "")},
        "lm_studio_online": bool(loaded) or lm_reachable(),
        "loaded_models": loaded,
    }


@app.get("/api/status")
def api_status(x_init_data: Optional[str] = Header(default=None)):
    resolve_identity(x_init_data)
    lm_online = lm_reachable()
    loaded = get_loaded_models() if lm_online else []
    db_exists = os.path.exists(RAG_DB_PATH)
    emb_loaded = _embedding_model is not None
    bot_ok = telegram_bot_configured()

    def comp(name, status, detail):
        return {"name": name, "status": status, "detail": detail}

    components = [
        comp("Mini App backend", "ok", "работает и отвечает"),
        comp("LM Studio", "ok" if lm_online else "err",
             (f"онлайн · моделей: {len(loaded)}" if lm_online else "не отвечает на :1234")),
        comp("RAG база", "ok" if db_exists else "err",
             (get_db_size_text() if db_exists else "файл не найден")),
        comp("Модель эмбеддингов", "ok" if emb_loaded else "warn",
             ("загружена" if emb_loaded else "ленивая загрузка при поиске")),
        comp("Telegram бот",
             "ok" if bot_ok else ("warn" if bot_ok is None else "err"),
             ("токен валиден" if bot_ok else ("токен не задан" if bot_ok is None else "токен недоступен"))),
        comp("Whisper (голос)", "ok" if _whisper_model is not None else "na",
             ("загружена" if _whisper_model is not None else "загрузится при первом голосовом")),
    ]
    overall = "ok"
    if any(c["status"] == "err" for c in components):
        overall = "err"
    elif any(c["status"] == "warn" for c in components):
        overall = "warn"
    return {"overall": overall, "components": components}


@app.get("/api/models")
def api_models(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    loaded = set(get_loaded_models())
    models = [
        {"key": k, "id": i["id"], "label": i["label"], "loaded": i["id"] in loaded}
        for k, i in AVAILABLE_MODELS.items()
    ]
    return {"models": models, "selected": webapp_selected_model.get(chat_id, DEFAULT_MODEL_KEY)}


def _background_load(chat_id, model_id, key):
    ok, err = ensure_model_loaded(model_id)
    # Отмечаем результат только если пользователь не переключился на другую модель.
    cur = model_load_status.get(chat_id)
    if cur and cur.get("key") == key:
        model_load_status[chat_id] = {"key": key, "state": "ready" if ok else "error", "error": err}


@app.post("/api/models/select")
def api_models_select(payload: dict = Body(...), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    key = payload.get("key")
    if key not in AVAILABLE_MODELS:
        raise HTTPException(status_code=400, detail="Неизвестная модель")
    webapp_selected_model[chat_id] = key
    if payload.get("load", True):
        # Запускаем загрузку в фоне и сразу отвечаем — интерфейс не висит.
        model_load_status[chat_id] = {"key": key, "state": "loading", "error": None}
        threading.Thread(
            target=_background_load,
            args=(chat_id, AVAILABLE_MODELS[key]["id"], key),
            daemon=True,
        ).start()
        return {"selected": key, "loading": True}
    return {"selected": key, "loading": False}


@app.get("/api/models/loadstatus")
def api_models_loadstatus(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    return model_load_status.get(chat_id) or {"state": "idle", "key": None, "error": None}


@app.get("/api/settings")
def api_get_settings(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    return get_settings(chat_id)


@app.post("/api/settings")
def api_set_settings(payload: dict = Body(...), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    key = payload.get("key")
    value = bool(payload.get("value"))
    if key not in {"rag_only", "study_mode"}:
        raise HTTPException(status_code=400, detail="Неизвестная настройка")
    return set_setting(chat_id, key, value)


@app.get("/api/documents")
def api_documents(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    total_chunks, docs = get_rag_stats(chat_id)
    return {"total_chunks": total_chunks, "documents": docs}


@app.get("/api/documents/{doc_id}")
def api_document_detail(doc_id: int, x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    doc = get_document_by_id(chat_id, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не найден")
    chunks = get_document_chunks(chat_id, doc_id, limit=5)
    doc["preview_chunks"] = [
        {"index": ci, "text": (t[:600] + "…") if len(t) > 600 else t} for ci, t in chunks
    ]
    return doc


@app.delete("/api/documents/{doc_id}")
def api_document_delete(doc_id: int, x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    if not delete_rag_document_by_id(chat_id, doc_id):
        raise HTTPException(status_code=404, detail="Документ не найден")
    return {"deleted": True}


@app.get("/api/documents/{doc_id}/download")
def api_document_download(doc_id: int, init_data: Optional[str] = Query(default=None)):
    chat_id, _, _ = resolve_identity(init_data)
    doc = get_document_by_id(chat_id, doc_id)
    if not doc or not doc.get("saved_path") or not os.path.exists(doc["saved_path"]):
        raise HTTPException(status_code=404, detail="Файл недоступен")
    return FileResponse(doc["saved_path"], filename=doc["source_name"], media_type="application/octet-stream")


@app.post("/api/upload")
async def api_upload(file: UploadFile = File(...), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    file_name = file.filename or "document"
    ext = os.path.splitext(file_name.lower())[1]
    if ext not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Формат не поддерживается (PDF, TXT, MD, LOG, CSV, JSON, DOCX)")
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Пустой файл")
    if len(file_bytes) > RAG_MAX_FILE_MB * 1024 * 1024:
        raise HTTPException(status_code=400, detail=f"Файл больше {RAG_MAX_FILE_MB} МБ")

    file_hash = hashlib.sha256(file_bytes).hexdigest()
    existing = get_document_by_hash(chat_id, file_hash)
    if existing:
        return {"status": "duplicate", "document": existing, "inserted": 0}

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        text = extract_text_from_file(tmp_path, file_name)
        saved_path = save_original_document(chat_id, file_name, file_bytes, file_hash)
        doc_id, inserted, total = add_document_to_rag(
            chat_id, file_name, text, source_hash=file_hash, saved_path=saved_path,
            file_size=len(file_bytes), file_ext=ext,
        )
        if not inserted:
            raise HTTPException(status_code=422, detail="Не удалось извлечь текст (пустой/скан)")
        return {"status": "ok", "document": get_document_by_id(chat_id, doc_id),
                "inserted": inserted, "chunks": total}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Ошибка обработки: {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


@app.post("/api/transcribe")
async def api_transcribe(file: UploadFile = File(...), x_init_data: Optional[str] = Header(default=None)):
    resolve_identity(x_init_data)
    audio = await file.read()
    if not audio:
        raise HTTPException(status_code=400, detail="Пустой аудиофайл")
    suffix = os.path.splitext(file.filename or "audio.webm")[1] or ".webm"
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(audio)
            tmp_path = tmp.name
        model = get_whisper()
        segments, info = model.transcribe(tmp_path, vad_filter=True)
        text = " ".join(s.text.strip() for s in segments).strip()
        return {"text": text, "language": getattr(info, "language", None)}
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Ошибка распознавания: {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


@app.get("/api/search")
def api_search(q: str = Query(..., min_length=2), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    try:
        _, sources = search_rag_with_sources(chat_id, q, top_k=8)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка поиска: {e}")
    results = [
        {
            "source_name": s["source_name"], "chunk_index": s["chunk_index"],
            "score": round(s["score"], 3), "semantic": round(s["semantic_score"], 3),
            "keyword": round(s["keyword_score"], 3),
            "text": (s["text"][:800] + "…") if len(s["text"]) > 800 else s["text"],
        }
        for s in sources
    ]
    return {"query": q, "results": results}


@app.get("/api/sessions")
def api_sessions(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    return {"sessions": list_sessions(chat_id)}


@app.post("/api/sessions")
def api_sessions_new(x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    sid = create_session(chat_id)
    return {"id": sid, "title": "Новый чат"}


@app.delete("/api/sessions/{session_id}")
def api_sessions_delete(session_id: int, x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    delete_session(chat_id, session_id)
    conversation_history.pop(f"{chat_id}:{session_id}", None)
    return {"deleted": True}


@app.get("/api/chat/history")
def api_chat_history(session: int = Query(...), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    return {"messages": load_session_messages(chat_id, session)}


@app.post("/api/chat")
def api_chat(payload: dict = Body(...), x_init_data: Optional[str] = Header(default=None)):
    chat_id, _, _ = resolve_identity(x_init_data)
    user_message = (payload.get("message") or "").strip()
    image = payload.get("image")  # data URL (base64) для vision-моделей
    if not user_message and not image:
        raise HTTPException(status_code=400, detail="Пустое сообщение")

    settings = get_settings(chat_id)
    rag_only = settings["rag_only"]
    study_mode = settings["study_mode"]
    model_key = webapp_selected_model.get(chat_id, DEFAULT_MODEL_KEY)
    model_id = AVAILABLE_MODELS[model_key]["id"]

    sid = payload.get("session_id") or create_session(chat_id)
    key = f"{chat_id}:{sid}"
    if key not in conversation_history:
        conversation_history[key] = [{"role": "system", "content": SYSTEM_PROMPT}]
        for m in load_session_messages(chat_id, sid):
            conversation_history[key].append({"role": m["role"], "content": m["content"]})
        while len(conversation_history[key]) > 12:
            if len(conversation_history[key]) > 1:
                conversation_history[key].pop(1)

    # Регенерация: убираем прошлую пару (вопрос+ответ), чтобы заменить ответ без дублей.
    if payload.get("regen"):
        for _ in range(2):
            hist = conversation_history.get(key, [])
            if len(hist) > 1 and hist[-1]["role"] in ("assistant", "user"):
                hist.pop()
        delete_last_messages(sid, 2)

    rag_context, used_sources = "", []
    try:
        rag_context, used_sources = search_rag_with_sources(chat_id, user_message)
    except Exception as e:
        print(f"RAG search error: {e}")

    def error_stream(text):
        yield json.dumps({"type": "error", "text": text}, ensure_ascii=False) + "\n"

    if rag_only and not rag_context:
        return StreamingResponse(
            error_stream("📚 Режим RAG-only включён. В загруженных документах ответа не нашлось."),
            media_type="application/x-ndjson",
        )

    messages = list(conversation_history[key])
    if study_mode:
        messages.append({"role": "system", "content": STUDY_MODE_PROMPT})
    if rag_context:
        messages.append({"role": "system", "content": build_rag_system_message(rag_context, rag_only)})
    if image:
        messages.append({"role": "user", "content": [
            {"type": "text", "text": user_message or "Опиши это изображение."},
            {"type": "image_url", "image_url": {"url": image}},
        ]})
    else:
        messages.append({"role": "user", "content": user_message})

    sources_meta, seen = [], set()
    for s in used_sources:
        # ВАЖНО: имя переменной не должно совпадать с `key` (chat_id:sid) из внешней области
        # видимости — иначе conversation_history[key] внутри stream() ломается при наличии RAG-источников.
        src_key = (s["source_name"], s["chunk_index"])
        if src_key in seen:
            continue
        seen.add(src_key)
        sources_meta.append({"source_name": s["source_name"], "chunk_index": s["chunk_index"] + 1})

    def stream():
        full_response = ""
        yield json.dumps({"type": "session", "id": sid}, ensure_ascii=False) + "\n"
        if sources_meta:
            yield json.dumps({"type": "sources", "sources": sources_meta}, ensure_ascii=False) + "\n"
        data = {"model": model_id, "messages": messages, "stream": True}
        try:
            with requests.post(LM_STUDIO_CHAT_URL, headers=get_headers(), json=data,
                               stream=True, timeout=300) as response:
                if response.status_code == 503:
                    yield json.dumps({"type": "error", "text": "⚠️ Модель недоступна. Запущен ли сервер LM Studio?"},
                                     ensure_ascii=False) + "\n"
                    return
                if response.status_code != 200:
                    yield json.dumps({"type": "error", "text": f"Ошибка API: {response.status_code}"},
                                     ensure_ascii=False) + "\n"
                    return
                for line in response.iter_lines():
                    if not line:
                        continue
                    line_text = line.decode("utf-8")
                    if not line_text.startswith("data: "):
                        continue
                    chunk_payload = line_text[len("data: "):]
                    if chunk_payload.strip() == "[DONE]":
                        break
                    try:
                        chunk = json.loads(chunk_payload)
                    except json.JSONDecodeError:
                        continue
                    delta = chunk.get("choices", [{}])[0].get("delta", {})
                    piece = delta.get("content", "")
                    if piece:
                        full_response += piece
                        yield json.dumps({"type": "token", "text": piece}, ensure_ascii=False) + "\n"
        except requests.exceptions.ConnectionError:
            yield json.dumps({"type": "error", "text": "⚠️ Сервер LM Studio недоступен."},
                             ensure_ascii=False) + "\n"
            return
        except Exception as e:
            yield json.dumps({"type": "error", "text": f"Ошибка: {e}"}, ensure_ascii=False) + "\n"
            return

        if full_response:
            stored_user = user_message or "🖼 [изображение]"
            conversation_history[key].append({"role": "user", "content": stored_user})
            conversation_history[key].append({"role": "assistant", "content": full_response})
            while len(conversation_history[key]) > 12:
                if len(conversation_history[key]) > 1:
                    conversation_history[key].pop(1)
            try:
                save_chat_message(chat_id, sid, "user", stored_user)
                save_chat_message(chat_id, sid, "assistant", full_response)
                set_session_title(sid, stored_user)
                touch_session(sid)
            except Exception as e:
                print(f"save history error: {e}")
        yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"

    return StreamingResponse(stream(), media_type="application/x-ndjson")


@app.get("/api/admin/stats")
def api_admin_stats(x_init_data: Optional[str] = Header(default=None)):
    _, is_admin, _ = resolve_identity(x_init_data)
    require_admin(is_admin)
    return get_admin_stats()


@app.get("/api/admin/users")
def api_admin_users(limit: int = Query(default=30, ge=1, le=200),
                    x_init_data: Optional[str] = Header(default=None)):
    _, is_admin, _ = resolve_identity(x_init_data)
    require_admin(is_admin)
    return {"users": get_admin_users(limit)}


# ------------------------------------------------------------------
# Фронтенд.
# ------------------------------------------------------------------
@app.get("/")
def index():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))


if os.path.isdir(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("MINIAPP_PORT", "8080"))
    print(f"Mini App backend: http://localhost:{port}")
    print(f"RAG DB: {RAG_DB_PATH}")
    uvicorn.run(app, host="0.0.0.0", port=port)
