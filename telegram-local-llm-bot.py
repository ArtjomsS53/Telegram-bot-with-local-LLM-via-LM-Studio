import os
import time
import json
import random
import base64
import requests
import sqlite3
import hashlib
import tempfile
import re
import shutil
import asyncio
import uuid
import zipfile
import glob
import html
import numpy as np
from dotenv import load_dotenv

# Папка проекта = папка, где лежит этот .py файл.
# Так база RAG и .env не будут зависеть от того, откуда IDE запускает Python.
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

# Загружаем .env именно из папки проекта.
load_dotenv(os.path.join(PROJECT_DIR, ".env"))


def project_path(path):
    """Возвращает абсолютный путь внутри папки проекта, если путь относительный."""
    if os.path.isabs(path):
        return path
    return os.path.join(PROJECT_DIR, path)

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from faster_whisper import WhisperModel

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None
from telegram.constants import ParseMode
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler, filters

# Истории диалогов и выбранные модели по chat_id
conversation_history = {}
user_selected_model = {}
pending_user_content = {}  # Сообщения, ожидающие выбора модели

TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")

LM_STUDIO_BASE_URL = "http://localhost:1234"
LM_STUDIO_CHAT_URL = f"{LM_STUDIO_BASE_URL}/v1/chat/completions"
LM_STUDIO_MODELS_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models"
LM_STUDIO_LOAD_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models/load"
LM_STUDIO_UNLOAD_URL = f"{LM_STUDIO_BASE_URL}/api/v1/models/unload"

SYSTEM_PROMPT = "Ты — полезный ассистент. Отвечай ТОЛЬКО на русском языке. Давай краткие, естественные ответы."

DEFAULT_MODEL_KEY = "balanced"

AVAILABLE_MODELS = {
    "fast": {"id": "google/gemma-4-e4b", "label": "⚡ Быстрая (117 т/с)"},
    "balanced": {"id": "google/gemma-4-12b-qat", "label": "⚖️ Сбалансированная (70 т/с)"},
    "smart": {"id": "google/gemma-4-26b-a4b-qat", "label": "🧠 Умная (33 т/с)"},
}

WHISPER_MODEL_NAME = "medium"

print(f"Загружаю модель распознавания речи (Whisper: {WHISPER_MODEL_NAME})...")
whisper_model = WhisperModel(WHISPER_MODEL_NAME, device="cpu", compute_type="int8")
print("Whisper модель загружена.")

THINKING_PHRASES = [
    "🤔 Думаю над ответом...",
    "💭 Обрабатываю запрос...",
    "⚡ Генерирую ответ...",
    "🧠 Анализирую вопрос...",
]

currently_loaded_model = None


# -------------------------
# RAG: память документов
# -------------------------
embedding_model = None

# Папка для всей RAG-памяти: оригинальные файлы + SQLite-база.
# По умолчанию база будет лежать тут:
#   <папка проекта>/rag_files/rag_memory.sqlite3
RAG_STORAGE_DIR = project_path(os.environ.get("RAG_STORAGE_DIR", "rag_files"))
RAG_DB_PATH = project_path(
    os.environ.get("RAG_DB_PATH", os.path.join(RAG_STORAGE_DIR, "rag_memory.sqlite3"))
)
RAG_EMBEDDING_MODEL_NAME = os.environ.get(
    "RAG_EMBEDDING_MODEL",
    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
)
RAG_TOP_K = int(os.environ.get("RAG_TOP_K", "4"))
RAG_MAX_CONTEXT_CHARS = int(os.environ.get("RAG_MAX_CONTEXT_CHARS", "4500"))
RAG_CHUNK_SIZE = int(os.environ.get("RAG_CHUNK_SIZE", "900"))
RAG_CHUNK_OVERLAP = int(os.environ.get("RAG_CHUNK_OVERLAP", "150"))
RAG_MIN_SCORE = float(os.environ.get("RAG_MIN_SCORE", "0.22"))
RAG_MAX_FILE_MB = int(os.environ.get("RAG_MAX_FILE_MB", "20"))
RAG_KEYWORD_WEIGHT = float(os.environ.get("RAG_KEYWORD_WEIGHT", "0.45"))
RAG_SEMANTIC_WEIGHT = float(os.environ.get("RAG_SEMANTIC_WEIGHT", "0.55"))
RAG_SEARCH_RESULTS = int(os.environ.get("RAG_SEARCH_RESULTS", "6"))

MAX_TELEGRAM_MESSAGE_CHARS = int(os.environ.get("MAX_TELEGRAM_MESSAGE_CHARS", "3800"))
TELEGRAM_RICH_TEXT = os.environ.get("TELEGRAM_RICH_TEXT", "true").strip().lower() in {"1", "true", "yes", "on"}
PLAIN_TEXT_OUTPUT = os.environ.get("PLAIN_TEXT_OUTPUT", "false").strip().lower() in {"1", "true", "yes", "on"}

# Если TELEGRAM_RICH_TEXT=true, бот отправляет ответы с безопасным HTML-форматированием.
# Если PLAIN_TEXT_OUTPUT=true, форматирование будет принудительно очищаться до обычного текста.

# Админка. В .env можно указать, например:
# ADMIN_TELEGRAM_IDS=123456789,987654321
# ADMIN_ONLY_MODE=false
ADMIN_TELEGRAM_IDS = {
    item.strip()
    for item in os.environ.get("ADMIN_TELEGRAM_IDS", "").split(",")
    if item.strip()
}
ADMIN_ONLY_MODE = os.environ.get("ADMIN_ONLY_MODE", "false").strip().lower() in {"1", "true", "yes", "on"}

SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".md", ".log", ".csv", ".json", ".docx"}



def migrate_old_rag_db_if_needed():
    """
    Если старая rag_memory.sqlite3 лежит в корне проекта или в рабочей папке запуска,
    один раз копируем её в rag_files/rag_memory.sqlite3, чтобы память не пропала.
    """
    try:
        new_path = os.path.abspath(RAG_DB_PATH)
        old_candidates = [
            os.path.join(PROJECT_DIR, "rag_memory.sqlite3"),
            os.path.abspath("rag_memory.sqlite3"),
        ]

        # убираем дубли, сохраняя порядок
        seen = set()
        old_candidates = [
            path for path in old_candidates
            if not (path in seen or seen.add(path))
        ]

        for old_path in old_candidates:
            old_path = os.path.abspath(old_path)

            if old_path == new_path:
                continue

            if os.path.exists(old_path) and not os.path.exists(new_path):
                os.makedirs(os.path.dirname(new_path), exist_ok=True)
                shutil.copy2(old_path, new_path)
                print(f"RAG: старая база скопирована в rag_files: {new_path}")
                return

    except Exception as e:
        print(f"RAG: не удалось перенести старую базу: {e}")

def init_rag_db():
    """Создаёт SQLite-базу для локальной памяти документов."""
    migrate_old_rag_db_if_needed()
    os.makedirs(os.path.dirname(os.path.abspath(RAG_DB_PATH)), exist_ok=True)

    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rag_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT NOT NULL,
                source_name TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                chunk_hash TEXT NOT NULL,
                text TEXT NOT NULL,
                embedding BLOB NOT NULL,
                created_at INTEGER NOT NULL,
                UNIQUE(chat_id, chunk_hash)
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_rag_chat_id ON rag_chunks(chat_id)")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rag_settings (
                chat_id TEXT PRIMARY KEY,
                rag_only INTEGER NOT NULL DEFAULT 0,
                updated_at INTEGER NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS bot_users (
                chat_id TEXT PRIMARY KEY,
                user_id TEXT,
                username TEXT,
                first_name TEXT,
                last_name TEXT,
                first_seen_at INTEGER NOT NULL,
                last_seen_at INTEGER NOT NULL,
                message_count INTEGER NOT NULL DEFAULT 0
            )
            """
        )
        conn.commit()


def get_embedding_model():
    """Ленивая загрузка модели embeddings, чтобы бот быстрее стартовал."""
    global embedding_model

    if SentenceTransformer is None:
        raise RuntimeError(
            "Не установлен sentence-transformers. Установи: pip install sentence-transformers"
        )

    if embedding_model is None:
        print(f"Загружаю embedding-модель для RAG: {RAG_EMBEDDING_MODEL_NAME}")
        embedding_model = SentenceTransformer(RAG_EMBEDDING_MODEL_NAME)
        print("Embedding-модель загружена.")

    return embedding_model


def embed_texts(texts):
    """Возвращает нормализованные embeddings в float32."""
    model = get_embedding_model()
    embeddings = model.encode(
        texts,
        normalize_embeddings=True,
        show_progress_bar=False
    )
    return np.asarray(embeddings, dtype=np.float32)


def clean_text(text):
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()



def clean_model_output(text):
    """
    Убирает Markdown-разметку из ответа модели, чтобы в Telegram не торчали **, #, >, ` и т.п.
    Это не делает ответ красивее через Telegram parse_mode, а именно превращает его в обычный текст.
    """
    if text is None:
        return ""

    t = str(text).replace("\r\n", "\n").replace("\r", "\n")

    # Убираем fenced-code блоки, оставляя содержимое.
    t = re.sub(r"```[a-zA-Z0-9_+\-.]*\n?", "", t)
    t = t.replace("```", "")

    # Заголовки Markdown: ### Заголовок -> Заголовок
    t = re.sub(r"(?m)^\s{0,3}#{1,6}\s*", "", t)

    # Цитаты Markdown: > текст -> текст
    t = re.sub(r"(?m)^\s{0,3}>\s?", "", t)

    # Горизонтальные линии Markdown.
    t = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", t)

    # Markdown-списки через * / - / + переводим в обычный маркер.
    t = re.sub(r"(?m)^\s*[\*\-\+]\s+", "• ", t)

    # Чекбоксы Markdown.
    t = re.sub(r"(?mi)^\s*[-*+]\s*\[[ x]\]\s*", "• ", t)

    # Ссылки Markdown: [текст](url) -> текст (url)
    t = re.sub(r"\[([^\]\n]+)\]\(([^)\n]+)\)", r"\1 (\2)", t)

    # Жирный / курсив / inline-code. Делаем несколько проходов для вложенных случаев.
    for _ in range(4):
        t = re.sub(r"\*\*([^*\n]+)\*\*", r"\1", t)
        t = re.sub(r"__([^_\n]+)__", r"\1", t)
        t = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", t)
        t = re.sub(r"(?<!_)_([^_\n]+)_(?!_)", r"\1", t)
        t = re.sub(r"`([^`\n]+)`", r"\1", t)

    # Остатки типичной разметки на краях строк.
    t = re.sub(r"(?m)^\s*[*_`#>]+\s*", "", t)
    t = t.replace("**", "").replace("__", "").replace("`", "")

    # Если модель сделала Markdown-таблицу, убираем разделители вида |---|---|.
    t = re.sub(r"(?m)^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", "", t)

    # Чистим лишние пустые строки и пробелы.
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def markdown_to_telegram_html(text):
    """
    Безопасно превращает простой Markdown от LLM в Telegram HTML.
    Telegram поддерживает ограниченный набор тегов, поэтому заголовки переводим в <b>,
    списки оставляем обычными строками с маркером, а весь обычный текст экранируем.
    """
    if text is None:
        return ""

    raw = str(text).replace("\r\n", "\n").replace("\r", "\n")
    placeholders = {}

    def put_placeholder(rendered_html):
        key = f"@@HTML_BLOCK_{len(placeholders)}@@"
        placeholders[key] = rendered_html
        return key

    # Fenced code blocks: ```python ... ``` -> <pre>...</pre>
    def code_block_repl(match):
        code = match.group(2) or ""
        return put_placeholder(f"<pre>{html.escape(code.strip())}</pre>")

    raw = re.sub(r"```([a-zA-Z0-9_+\-.]*)\n([\s\S]*?)```", code_block_repl, raw)

    # Inline code: `text` -> <code>text</code>
    def inline_code_repl(match):
        code = match.group(1) or ""
        return put_placeholder(f"<code>{html.escape(code)}</code>")

    raw = re.sub(r"`([^`\n]+)`", inline_code_repl, raw)

    # Экранируем всё, что не стало placeholder, чтобы Telegram не падал на <, >, &.
    t = html.escape(raw)

    # Markdown links: [text](https://...) -> <a href="...">text</a>
    def link_repl(match):
        label = match.group(1)
        url = html.unescape(match.group(2)).strip()
        if not re.match(r"(?i)^(https?://|tg://|mailto:)", url):
            return f"{label} ({html.escape(url)})"
        safe_url = html.escape(url, quote=True)
        return f'<a href="{safe_url}">{label}</a>'

    t = re.sub(r"\[([^\]\n]+)\]\(([^)\n]+)\)", link_repl, t)

    # Заголовки Markdown: # Текст -> <b>Текст</b>
    t = re.sub(r"(?m)^\s{0,3}#{1,6}\s+(.+?)\s*$", r"<b>\1</b>", t)

    # Цитаты Markdown. Telegram HTML поддерживает blockquote в актуальном Bot API,
    # но fallback ниже всё равно отправит plain text, если клиент/API ругнётся.
    t = re.sub(r"(?m)^\s{0,3}&gt;\s?(.+?)\s*$", r"<blockquote>\1</blockquote>", t)

    # Горизонтальные линии Markdown убираем.
    t = re.sub(r"(?m)^\s*[-*_]{3,}\s*$", "", t)

    # Чекбоксы и списки: Telegram не поддерживает <ul>/<li>, поэтому делаем красивый текстовый маркер.
    t = re.sub(r"(?mi)^\s*[-*+]\s*\[[ x]\]\s*", "• ", t)
    t = re.sub(r"(?m)^\s*[-*+]\s+", "• ", t)

    # Жирный, курсив, зачёркивание. Несколько проходов для простых вложенных случаев.
    for _ in range(3):
        t = re.sub(r"\*\*([^*\n]+?)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"__([^_\n]+?)__", r"<b>\1</b>", t)
        t = re.sub(r"~~([^~\n]+?)~~", r"<s>\1</s>", t)
        t = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"<i>\1</i>", t)
        t = re.sub(r"(?<!_)_([^_\n]+?)_(?!_)", r"<i>\1</i>", t)

    # Остатки частой разметки, чтобы не торчали ** или ``` при плохом ответе модели.
    t = t.replace("```", "").replace("**", "").replace("__", "")
    t = re.sub(r"(?m)^\s{0,3}#{1,6}\s*", "", t)

    # Убираем markdown-разделители таблиц вида |---|---|, сами строки таблицы оставляем.
    t = re.sub(r"(?m)^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", "", t)

    # Возвращаем code placeholders уже после обработки Markdown.
    for key, rendered in placeholders.items():
        t = t.replace(key, rendered)

    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

def split_text_into_chunks(text, chunk_size=RAG_CHUNK_SIZE, overlap=RAG_CHUNK_OVERLAP):
    """Режет текст на куски с небольшим overlap, чтобы не терять смысл на границах."""
    text = clean_text(text)
    if not text:
        return []

    compact = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0

    while start < len(compact):
        end = min(len(compact), start + chunk_size)

        if end < len(compact):
            # Пытаемся закончить кусок на границе предложения, а не посреди слова.
            sentence_end = max(
                compact.rfind(". ", start, end),
                compact.rfind("! ", start, end),
                compact.rfind("? ", start, end),
                compact.rfind("; ", start, end),
            )
            if sentence_end > start + int(chunk_size * 0.55):
                end = sentence_end + 1

        chunk = compact[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(compact):
            break

        start = max(0, end - overlap)
        if start >= end:
            start = end

    return chunks


def extract_text_from_pdf(file_path):
    """
    Пытается вытащить текст из PDF.
    Сначала через pypdf, если он падает — через PyMuPDF.
    """
    errors = []

    # Способ 1: pypdf
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path, strict=False)
        pages_text = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                pages_text.append(page.extract_text() or "")
            except Exception as page_error:
                errors.append(f"pypdf page {page_number}: {page_error}")

        text = "\n".join(pages_text).strip()
        if text:
            return text

    except Exception as e:
        errors.append(f"pypdf: {e}")

    # Способ 2: PyMuPDF / fitz
    try:
        import fitz

        doc = fitz.open(file_path)
        pages_text = []

        for page_number in range(len(doc)):
            try:
                page = doc.load_page(page_number)
                pages_text.append(page.get_text("text") or "")
            except Exception as page_error:
                errors.append(f"fitz page {page_number + 1}: {page_error}")

        doc.close()

        text = "\n".join(pages_text).strip()
        if text:
            return text

    except Exception as e:
        errors.append(f"fitz: {e}")

    raise RuntimeError(
        "Не удалось извлечь текст из PDF. "
        "Возможно, это скан/картинка, битый PDF или PDF без текстового слоя. "
        f"Ошибки: {' | '.join(errors[:3])}"
    )


def extract_text_from_docx(file_path):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Не установлен python-docx. Установи: pip install python-docx")

    document = Document(file_path)
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Читаем ещё и таблицы, если они есть.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_from_plain_file(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    return raw.decode("utf-8", errors="ignore")


def extract_text_from_file(file_path, file_name):
    extension = os.path.splitext(file_name.lower())[1]

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)

    if extension == ".docx":
        return extract_text_from_docx(file_path)

    if extension in {".txt", ".md", ".log", ".csv", ".json"}:
        return extract_text_from_plain_file(file_path)

    raise ValueError(
        "Неподдерживаемый формат. Можно отправлять: PDF, TXT, MD, LOG, CSV, JSON, DOCX."
    )


def add_text_to_rag(chat_id, source_name, text):
    """Добавляет текст в RAG-память конкретного Telegram-чата."""
    init_rag_db()

    text = clean_text(text)
    chunks = split_text_into_chunks(text)

    if not chunks:
        return 0

    embeddings = embed_texts(chunks)
    now = int(time.time())
    inserted_count = 0

    with sqlite3.connect(RAG_DB_PATH) as conn:
        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_hash = hashlib.sha256(chunk.encode("utf-8")).hexdigest()
            try:
                conn.execute(
                    """
                    INSERT INTO rag_chunks
                    (chat_id, source_name, chunk_index, chunk_hash, text, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        str(chat_id),
                        source_name,
                        index,
                        chunk_hash,
                        chunk,
                        embedding.astype(np.float32).tobytes(),
                        now,
                    )
                )
                inserted_count += 1
            except sqlite3.IntegrityError:
                # Такой же кусок уже есть в памяти этого чата.
                pass
        conn.commit()

    return inserted_count



def extract_keywords(query):
    """Достаёт ключевые слова/коды из запроса для keyword-части гибридного поиска."""
    if not isinstance(query, str):
        return []

    tokens = re.findall(r"[a-zA-Zа-яА-ЯёЁ0-9_\-]{3,}", query.lower())
    stop_words = {
        "что", "как", "где", "когда", "какой", "какая", "какое", "какие",
        "это", "тот", "там", "его", "её", "она", "они", "или", "для", "про",
        "мне", "тебе", "надо", "нужно", "есть", "был", "была", "были",
        "the", "and", "for", "with", "what", "where", "when", "how", "who",
    }

    result = []
    seen = set()
    for token in tokens:
        if token in stop_words:
            continue
        if token not in seen:
            seen.add(token)
            result.append(token)
    return result


def calculate_keyword_score(query, text):
    """Оценка точных совпадений. Помогает находить коды, номера, имена и суммы."""
    query = (query or "").strip().lower()
    text_lower = (text or "").lower()

    if not query or not text_lower:
        return 0.0

    keywords = extract_keywords(query)
    if not keywords:
        return 0.0

    matches = 0
    important_matches = 0

    for token in keywords:
        if token in text_lower:
            matches += 1
            # Коды/номера с цифрами, дефисами и подчёркиваниями важнее обычных слов.
            if any(ch.isdigit() for ch in token) or "-" in token or "_" in token:
                important_matches += 1

    token_score = matches / max(1, len(keywords))
    important_boost = min(0.25, important_matches * 0.08)

    phrase_boost = 0.0
    if len(query) >= 10 and query in text_lower:
        phrase_boost = 0.35

    return min(1.0, token_score * 0.75 + important_boost + phrase_boost)


def rank_rag_chunks(chat_id, query, top_k=RAG_TOP_K, include_below_min=False):
    """
    Гибридный поиск по RAG:
    - semantic score через embeddings
    - keyword score по точным словам/кодам
    Возвращает список dict с комбинированной оценкой.
    """
    if not isinstance(query, str) or len(query.strip()) < 3:
        return []

    init_rag_db()

    with sqlite3.connect(RAG_DB_PATH) as conn:
        rows = conn.execute(
            "SELECT source_name, chunk_index, text, embedding FROM rag_chunks WHERE chat_id = ?",
            (str(chat_id),)
        ).fetchall()

    if not rows:
        return []

    query_embedding = embed_texts([query])[0]
    ranked = []

    for source_name, chunk_index, text, embedding_blob in rows:
        chunk_embedding = np.frombuffer(embedding_blob, dtype=np.float32)
        if chunk_embedding.size != query_embedding.size:
            semantic_score = 0.0
        else:
            semantic_score = float(np.dot(query_embedding, chunk_embedding))

        keyword_score = calculate_keyword_score(query, text)
        hybrid_score = (semantic_score * RAG_SEMANTIC_WEIGHT) + (keyword_score * RAG_KEYWORD_WEIGHT)

        # Если keyword очень сильный, поднимаем фрагмент даже при слабом embedding.
        if keyword_score >= 0.65:
            hybrid_score = max(hybrid_score, keyword_score)

        if include_below_min or semantic_score >= RAG_MIN_SCORE or keyword_score >= 0.35 or hybrid_score >= RAG_MIN_SCORE:
            ranked.append(
                {
                    "score": hybrid_score,
                    "semantic_score": semantic_score,
                    "keyword_score": keyword_score,
                    "source_name": source_name,
                    "chunk_index": chunk_index,
                    "text": text,
                }
            )

    ranked.sort(key=lambda item: item["score"], reverse=True)
    return ranked[:top_k]


def search_rag_with_sources(chat_id, query, top_k=RAG_TOP_K):
    """
    Ищет релевантные куски документов для вопроса пользователя.
    Возвращает:
    - rag_context: текстовый контекст для модели
    - used_sources: список источников, чтобы показать пользователю, откуда взят ответ
    """
    selected = rank_rag_chunks(chat_id, query, top_k=top_k, include_below_min=False)

    if not selected:
        return "", []

    context_parts = []
    used_sources = []
    total_chars = 0

    for item in selected:
        source_name = item["source_name"]
        chunk_index = item["chunk_index"]
        text = item["text"]
        score = item["score"]
        semantic_score = item["semantic_score"]
        keyword_score = item["keyword_score"]

        block = (
            f"[Источник: {source_name}, фрагмент {chunk_index + 1}, "
            f"общая релевантность {score:.2f}, "
            f"semantic {semantic_score:.2f}, keyword {keyword_score:.2f}]\n{text}"
        )

        if total_chars + len(block) > RAG_MAX_CONTEXT_CHARS:
            remaining = RAG_MAX_CONTEXT_CHARS - total_chars
            if remaining <= 300:
                break
            block = block[:remaining].rstrip() + "..."

        context_parts.append(block)
        used_sources.append(item)
        total_chars += len(block)

    return "\n\n---\n\n".join(context_parts), used_sources


def search_rag(chat_id, query, top_k=RAG_TOP_K):
    """Старый совместимый вызов: возвращает только текст контекста."""
    rag_context, _ = search_rag_with_sources(chat_id, query, top_k=top_k)
    return rag_context


def format_rag_used_sources(used_sources, max_sources=5):
    """Форматирует список источников, которые были подмешаны в ответ."""
    if not used_sources:
        return ""

    unique_sources = []
    seen = set()

    for source in used_sources:
        key = (source["source_name"], source["chunk_index"])
        if key in seen:
            continue

        seen.add(key)
        unique_sources.append(source)

        if len(unique_sources) >= max_sources:
            break

    if not unique_sources:
        return ""

    lines = ["", "📚 Использовал память:"]
    for index, source in enumerate(unique_sources, start=1):
        lines.append(
            f"{index}. {source['source_name']}, "
            f"фрагмент {source['chunk_index'] + 1} "
            f"(общая {source['score']:.2f}, "
            f"semantic {source['semantic_score']:.2f}, "
            f"keyword {source['keyword_score']:.2f})"
        )

    return "\n".join(lines)


def build_rag_system_message(rag_context, rag_only=False):
    if rag_only:
        return (
            "Ты работаешь в режиме RAG-only. Отвечай ТОЛЬКО на основании фрагментов документов ниже.\n"
            "Если в этих фрагментах нет ответа, прямо скажи: 'В загруженных документах я не нашёл ответа.'\n"
            "Не используй внешние знания и не додумывай факты.\n\n"
            f"КОНТЕКСТ ИЗ ДОКУМЕНТОВ:\n{rag_context}"
        )

    return (
        "Ниже есть релевантные фрагменты из документов, которые пользователь ранее загрузил в память.\n"
        "Используй их только если они действительно помогают ответить на текущий вопрос.\n"
        "Если в документах нет ответа или информации мало, честно скажи об этом.\n"
        "Не выдумывай факты сверх предоставленного контекста.\n\n"
        f"КОНТЕКСТ ИЗ ДОКУМЕНТОВ:\n{rag_context}"
    )


def get_rag_stats(chat_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        total_chunks = conn.execute(
            "SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ?",
            (str(chat_id),)
        ).fetchone()[0]
        sources = conn.execute(
            """
            SELECT source_name, COUNT(*)
            FROM rag_chunks
            WHERE chat_id = ?
            GROUP BY source_name
            ORDER BY MAX(created_at) DESC
            """,
            (str(chat_id),)
        ).fetchall()
    return total_chunks, sources


def get_rag_documents(chat_id):
    """Возвращает список документов в RAG-памяти этого чата."""
    init_rag_db()

    with sqlite3.connect(RAG_DB_PATH) as conn:
        rows = conn.execute(
            """
            SELECT
                source_name,
                COUNT(*) AS chunks_count,
                MIN(created_at) AS first_added_at,
                MAX(created_at) AS last_added_at
            FROM rag_chunks
            WHERE chat_id = ?
            GROUP BY source_name
            ORDER BY last_added_at DESC
            """,
            (str(chat_id),)
        ).fetchall()

    return rows


def delete_rag_document(chat_id, source_name):
    """
    Удаляет конкретный документ из RAG-памяти.
    Сначала ищет точное совпадение, потом совпадение без учёта регистра.
    """
    init_rag_db()
    source_name = (source_name or "").strip()

    if not source_name:
        return 0, None

    documents = get_rag_documents(chat_id)

    matched_name = None
    for doc_name, *_ in documents:
        if doc_name == source_name:
            matched_name = doc_name
            break

    if matched_name is None:
        matches = [
            doc_name
            for doc_name, *_ in documents
            if doc_name.casefold() == source_name.casefold()
        ]

        if len(matches) == 1:
            matched_name = matches[0]

    if matched_name is None:
        return 0, None

    with sqlite3.connect(RAG_DB_PATH) as conn:
        cursor = conn.execute(
            "DELETE FROM rag_chunks WHERE chat_id = ? AND source_name = ?",
            (str(chat_id), matched_name)
        )
        conn.commit()
        deleted_count = cursor.rowcount

    return deleted_count, matched_name



def delete_rag_document_by_id(chat_id, doc_id):
    init_rag_db()
    doc = get_document_by_id(chat_id, int(doc_id))
    if not doc:
        return 0, None
    with sqlite3.connect(RAG_DB_PATH) as conn:
        cursor = conn.execute("DELETE FROM rag_chunks WHERE chat_id = ? AND document_id = ?", (str(chat_id), int(doc_id)))
        deleted_count = cursor.rowcount
        conn.execute("DELETE FROM rag_documents WHERE chat_id = ? AND id = ?", (str(chat_id), int(doc_id)))
        conn.commit()
    saved_path = doc.get("saved_path") or ""
    if saved_path and os.path.exists(saved_path):
        try:
            os.remove(saved_path)
        except Exception:
            pass
    return deleted_count, doc["source_name"]


def clear_rag_memory(chat_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        cursor = conn.execute("DELETE FROM rag_chunks WHERE chat_id = ?", (str(chat_id),))
        conn.commit()
        return cursor.rowcount




def get_rag_only_mode(chat_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        row = conn.execute(
            "SELECT rag_only FROM rag_settings WHERE chat_id = ?",
            (str(chat_id),)
        ).fetchone()
    return bool(row and row[0] == 1)


def set_rag_only_mode(chat_id, enabled):
    init_rag_db()
    now = int(time.time())
    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            """
            INSERT INTO rag_settings (chat_id, rag_only, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(chat_id) DO UPDATE SET
                rag_only = excluded.rag_only,
                updated_at = excluded.updated_at
            """,
            (str(chat_id), 1 if enabled else 0, now)
        )
        conn.commit()


def remember_user_in_db(update):
    """Запоминает базовую статистику по пользователю для админ-команд."""
    try:
        if not update.effective_chat:
            return

        chat_id = str(update.effective_chat.id)
        user = update.effective_user
        user_id = str(user.id) if user else ""
        username = user.username if user and user.username else ""
        first_name = user.first_name if user and user.first_name else ""
        last_name = user.last_name if user and user.last_name else ""
        now = int(time.time())

        init_rag_db()
        with sqlite3.connect(RAG_DB_PATH) as conn:
            conn.execute(
                """
                INSERT INTO bot_users
                (chat_id, user_id, username, first_name, last_name, first_seen_at, last_seen_at, message_count)
                VALUES (?, ?, ?, ?, ?, ?, ?, 1)
                ON CONFLICT(chat_id) DO UPDATE SET
                    user_id = excluded.user_id,
                    username = excluded.username,
                    first_name = excluded.first_name,
                    last_name = excluded.last_name,
                    last_seen_at = excluded.last_seen_at,
                    message_count = bot_users.message_count + 1
                """,
                (chat_id, user_id, username, first_name, last_name, now, now)
            )
            conn.commit()
    except Exception as e:
        print(f"Не удалось обновить bot_users: {e}")


def get_effective_user_id(update):
    user = update.effective_user
    return str(user.id) if user else ""


def is_admin(update):
    user_id = get_effective_user_id(update)
    return bool(user_id and user_id in ADMIN_TELEGRAM_IDS)


async def check_bot_access(update):
    remember_user_in_db(update)

    if not ADMIN_ONLY_MODE:
        return True

    if is_admin(update):
        return True

    if update.effective_message:
        await update.effective_message.reply_text(
            "⛔ Бот сейчас в закрытом режиме.\n"
            f"Твой Telegram ID: {get_effective_user_id(update)}\n\n"
            "Чтобы открыть доступ, добавь этот ID в ADMIN_TELEGRAM_IDS в .env."
        )
    return False


async def require_admin(update):
    remember_user_in_db(update)

    if not ADMIN_TELEGRAM_IDS:
        if update.effective_message:
            await update.effective_message.reply_text(
                "⚠️ Админка не настроена.\n\n"
                "Добавь в .env строку:\n"
                "ADMIN_TELEGRAM_IDS=твой_telegram_id\n\n"
                f"Твой Telegram ID: {get_effective_user_id(update)}"
            )
        return False

    if not is_admin(update):
        if update.effective_message:
            await update.effective_message.reply_text("⛔ Эта команда доступна только администратору.")
        return False

    return True


def get_db_size_text():
    if not os.path.exists(RAG_DB_PATH):
        return "0 Б"
    size = os.path.getsize(RAG_DB_PATH)
    if size < 1024:
        return f"{size} Б"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} КБ"
    return f"{size / (1024 * 1024):.2f} МБ"


def get_admin_stats():
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        users_count = conn.execute("SELECT COUNT(*) FROM bot_users").fetchone()[0]
        chats_with_chunks = conn.execute("SELECT COUNT(DISTINCT chat_id) FROM rag_chunks").fetchone()[0]
        total_chunks = conn.execute("SELECT COUNT(*) FROM rag_chunks").fetchone()[0]
        total_sources = conn.execute("SELECT COUNT(*) FROM (SELECT chat_id, source_name FROM rag_chunks GROUP BY chat_id, source_name)").fetchone()[0]
        rag_only_count = conn.execute("SELECT COUNT(*) FROM rag_settings WHERE rag_only = 1").fetchone()[0]

    return {
        "users_count": users_count,
        "chats_with_chunks": chats_with_chunks,
        "total_chunks": total_chunks,
        "total_sources": total_sources,
        "rag_only_count": rag_only_count,
        "db_size": get_db_size_text(),
        "db_path": RAG_DB_PATH,
    }


def get_admin_users(limit=20):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        return conn.execute(
            """
            SELECT chat_id, user_id, username, first_name, last_name, last_seen_at, message_count
            FROM bot_users
            ORDER BY last_seen_at DESC
            LIMIT ?
            """,
            (limit,)
        ).fetchall()


def get_headers():
    headers = {"Content-Type": "application/json"}
    if os.environ.get("LM_STUDIO_API_KEY"):
        headers["Authorization"] = f"Bearer {os.environ.get('LM_STUDIO_API_KEY')}"
    return headers


def get_loaded_models():
    """Возвращает список id моделей, которые сейчас реально загружены в LM Studio."""
    try:
        response = requests.get(LM_STUDIO_MODELS_URL, headers=get_headers(), timeout=15)
        if response.status_code != 200:
            return []

        data = response.json()
        loaded = []
        for item in data.get("models", []):
            instances = item.get("loaded_instances", [])
            for instance in instances:
                loaded.append(instance.get("id", item.get("key")))
        return loaded
    except Exception as e:
        print(f"Не удалось получить список загруженных моделей: {e}")
        return []


def unload_model(model_id):
    try:
        requests.post(
            LM_STUDIO_UNLOAD_URL,
            headers=get_headers(),
            json={"instance_id": model_id},
            timeout=30
        )
    except Exception as e:
        print(f"Не удалось выгрузить модель {model_id}: {e}")


def sync_loaded_models_on_startup():
    """При старте бота смотрит, что реально загружено, и оставляет только одну модель."""
    global currently_loaded_model

    known_model_ids = {info["id"] for info in AVAILABLE_MODELS.values()}
    loaded = get_loaded_models()
    loaded_known = [m for m in loaded if m in known_model_ids]

    if not loaded_known:
        print("При старте: ни одна известная модель не загружена.")
        currently_loaded_model = None
        return

    default_model_id = AVAILABLE_MODELS[DEFAULT_MODEL_KEY]["id"]
    keep = default_model_id if default_model_id in loaded_known else loaded_known[0]

    for model_id in loaded_known:
        if model_id != keep:
            print(f"Выгружаю лишнюю модель при старте: {model_id}")
            unload_model(model_id)

    currently_loaded_model = keep
    print(f"При старте оставлена загруженной модель: {keep}")


def trim_history(chat_id, max_messages=10):
    while len(conversation_history[chat_id]) > max_messages:
        if len(conversation_history[chat_id]) > 1:
            conversation_history[chat_id].pop(1)


def get_user_model_id(chat_id):
    model_key = user_selected_model.get(chat_id, DEFAULT_MODEL_KEY)
    return AVAILABLE_MODELS[model_key]["id"]


def ensure_model_loaded(model_id):
    global currently_loaded_model

    if currently_loaded_model == model_id:
        return True, None

    try:
        if currently_loaded_model:
            unload_model(currently_loaded_model)

        response = requests.post(
            LM_STUDIO_LOAD_URL,
            headers=get_headers(),
            json={"model": model_id},
            timeout=120
        )

        if response.status_code != 200:
            return False, f"Не удалось загрузить модель: {response.status_code}"

        currently_loaded_model = model_id
        return True, None

    except requests.exceptions.ConnectionError:
        return False, "Сервер LM Studio недоступен."
    except Exception as e:
        return False, str(e)


def build_model_keyboard(current_key=None):
    buttons = []
    for key, info in AVAILABLE_MODELS.items():
        label = info["label"]
        if key == current_key:
            label = f"✅ {label}"
        buttons.append([InlineKeyboardButton(label, callback_data=f"model:{key}")])
    return InlineKeyboardMarkup(buttons)



def split_telegram_text(text, max_chars=MAX_TELEGRAM_MESSAGE_CHARS):
    """Режет длинный текст на части, чтобы Telegram не падал на лимите 4096 символов."""
    text = str(text or "")
    if len(text) <= max_chars:
        return [text]

    chunks = []
    current = ""

    for line in text.splitlines(keepends=True):
        if len(line) > max_chars:
            if current:
                chunks.append(current.rstrip())
                current = ""
            for i in range(0, len(line), max_chars):
                chunks.append(line[i:i + max_chars].rstrip())
            continue

        if len(current) + len(line) > max_chars:
            chunks.append(current.rstrip())
            current = line
        else:
            current += line

    if current.strip():
        chunks.append(current.rstrip())

    return chunks or [""]


def render_for_telegram(text):
    """Готовит текст к отправке: либо plain text, либо безопасный Telegram HTML."""
    if PLAIN_TEXT_OUTPUT or not TELEGRAM_RICH_TEXT:
        return clean_model_output(text), None
    return markdown_to_telegram_html(text), ParseMode.HTML


async def safe_reply_long(message, text):
    """Безопасно отправляет длинный текст несколькими сообщениями, с HTML-форматированием и fallback."""
    raw_chunks = split_telegram_text(text)
    sent_messages = []
    for raw_chunk in raw_chunks:
        rendered, parse_mode = render_for_telegram(raw_chunk)
        try:
            sent_messages.append(await message.reply_text(rendered, parse_mode=parse_mode, disable_web_page_preview=True))
        except Exception:
            fallback = clean_model_output(raw_chunk)
            sent_messages.append(await message.reply_text(fallback[:MAX_TELEGRAM_MESSAGE_CHARS]))
    return sent_messages


async def safe_edit(thinking_msg, text):
    """Безопасно редактирует сообщение, сохраняя жирный текст/заголовки/код через Telegram HTML."""
    raw_chunks = split_telegram_text(text)
    first_raw = raw_chunks[0]
    rendered, parse_mode = render_for_telegram(first_raw)

    try:
        await thinking_msg.edit_text(rendered, parse_mode=parse_mode, disable_web_page_preview=True)
    except Exception:
        try:
            fallback = clean_model_output(first_raw)
            await thinking_msg.edit_text(fallback[:MAX_TELEGRAM_MESSAGE_CHARS])
        except Exception as e:
            if "Message is not modified" not in str(e):
                raise

    for raw_chunk in raw_chunks[1:]:
        rendered, parse_mode = render_for_telegram(raw_chunk)
        try:
            await thinking_msg.reply_text(rendered, parse_mode=parse_mode, disable_web_page_preview=True)
        except Exception:
            fallback = clean_model_output(raw_chunk)
            await thinking_msg.reply_text(fallback[:MAX_TELEGRAM_MESSAGE_CHARS])


async def ask_model_choice(update: Update, context, reason="restart"):
    """Показывает кнопки выбора модели с пояснением."""
    effective_msg = update.effective_message
    if reason == "restart":
        text = (
            "🔧 Бот был перезапущен разработчиком, поэтому активная модель сброшена.\n\n"
            "Пожалуйста, выберите, какую модель использовать дальше:"
        )
    else:
        text = "🎛 Выберите модель:"

    await effective_msg.reply_text(text, reply_markup=build_model_keyboard())


async def send_to_model(update: Update, context, user_content):
    effective_msg = update.effective_message
    chat_id = effective_msg.chat_id
    thinking_msg = None

    if not await check_bot_access(update):
        return

    model_id = get_user_model_id(chat_id)

    try:
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        thinking_msg = await effective_msg.reply_text(random.choice(THINKING_PHRASES))

        if chat_id not in conversation_history:
            conversation_history[chat_id] = [{"role": "system", "content": SYSTEM_PROMPT}]

        rag_context = ""
        rag_used_sources = []
        rag_only = get_rag_only_mode(chat_id)

        if isinstance(user_content, str):
            try:
                rag_context, rag_used_sources = search_rag_with_sources(chat_id, user_content)
            except Exception as rag_error:
                print(f"RAG search error: {rag_error}")
                rag_context = ""
                rag_used_sources = []

        if rag_only and not isinstance(user_content, str):
            await safe_edit(
                thinking_msg,
                "📚 Включён режим RAG-only. В этом режиме я отвечаю только на текстовые вопросы по загруженным документам."
            )
            return

        if rag_only and not rag_context:
            answer = (
                "📚 RAG-only включён.\n"
                "В загруженных документах я не нашёл ответа на этот вопрос."
            )
            await safe_edit(thinking_msg, answer)
            return

        ok, error = ensure_model_loaded(model_id)
        if not ok:
            await safe_edit(thinking_msg, f"⚠️ Проблема с моделью: {error}")
            return

        # Историю храним без RAG-контекста, чтобы она не раздувалась.
        # RAG-контекст добавляется только в конкретный запрос к модели.
        messages_for_model = list(conversation_history[chat_id])
        if rag_context:
            messages_for_model.append({"role": "system", "content": build_rag_system_message(rag_context, rag_only=rag_only)})

        messages_for_model.append({"role": "user", "content": user_content})

        data = {
            "model": model_id,
            "messages": messages_for_model,
            "stream": True
        }

        full_response = ""
        reasoning_buffer = ""
        token_count = 0
        start_time = time.time()
        last_update_time = start_time

        with requests.post(LM_STUDIO_CHAT_URL, headers=get_headers(), json=data, stream=True, timeout=180) as response:

            if response.status_code == 503:
                await safe_edit(thinking_msg, "⚠️ Локальная модель сейчас недоступна. Проверьте, запущен ли сервер LM Studio.")
                return

            if response.status_code != 200:
                raise Exception(f"API error: {response.status_code}")

            for line in response.iter_lines():
                if not line:
                    continue

                line_text = line.decode("utf-8")
                if not line_text.startswith("data: "):
                    continue

                payload = line_text[len("data: "):]
                if payload.strip() == "[DONE]":
                    break

                try:
                    chunk = json.loads(payload)
                except json.JSONDecodeError:
                    continue

                delta = chunk.get("choices", [{}])[0].get("delta", {})
                content_piece = delta.get("content", "")
                reasoning_piece = delta.get("reasoning_content", "")

                if content_piece:
                    full_response += content_piece
                    token_count += max(1, len(content_piece) // 4)

                if reasoning_piece:
                    reasoning_buffer += reasoning_piece
                    token_count += max(1, len(reasoning_piece) // 4)

                now = time.time()
                if now - last_update_time >= 2.5:
                    elapsed = now - start_time
                    speed = token_count / elapsed if elapsed > 0 else 0

                    if not full_response and reasoning_buffer:
                        status_text = f"🧠 Размышляю... (~{token_count} токенов, {speed:.0f} т/с)"
                    else:
                        if rag_context:
                            mode = "RAG-only" if rag_only else "RAG"
                            status_text = f"📚 Нашёл контекст ({mode}). Генерирую... (~{token_count} токенов, {speed:.0f} т/с)"
                        else:
                            status_text = f"⚡ Генерирую ответ... (~{token_count} токенов, {speed:.0f} т/с)"

                    try:
                        await thinking_msg.edit_text(status_text)
                    except Exception as edit_error:
                        if "Message is not modified" not in str(edit_error):
                            raise
                    last_update_time = now

        if not full_response:
            full_response = "🤷 Модель не смогла сформировать ответ, попробуйте переформулировать вопрос."

        if PLAIN_TEXT_OUTPUT:
            full_response = clean_model_output(full_response)

        sources_block = format_rag_used_sources(rag_used_sources)
        if sources_block:
            full_response = f"{full_response}\n{sources_block}"

        conversation_history[chat_id].append({"role": "user", "content": user_content})
        conversation_history[chat_id].append({"role": "assistant", "content": full_response})
        trim_history(chat_id)

        await safe_edit(thinking_msg, full_response)

    except requests.exceptions.ConnectionError:
        if thinking_msg:
            await safe_edit(thinking_msg, "⚠️ Локальная модель сейчас недоступна. Проверьте, запущен ли сервер LM Studio.")
        else:
            await effective_msg.reply_text("⚠️ Локальная модель сейчас недоступна.")

    except Exception as e:
        if thinking_msg:
            await safe_edit(thinking_msg, f"❌ Ошибка: {str(e)}")
        else:
            await effective_msg.reply_text(f"❌ Ошибка: {str(e)}")


async def remember_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    text = effective_msg.text or ""
    parts = text.split(maxsplit=1)

    if len(parts) < 2 or not parts[1].strip():
        await effective_msg.reply_text(
            "Напиши так:\n"
            "/remember текст, который надо запомнить\n\n"
            "Или просто отправь PDF/TXT/DOCX файлом — я сам добавлю его в память."
        )
        return

    try:
        status_msg = await effective_msg.reply_text("🧠 Добавляю текст в RAG-память...")
        inserted_count = add_text_to_rag(chat_id, "Текст из /remember", parts[1].strip())

        if inserted_count == 0:
            await status_msg.edit_text("⚠️ Не добавил новых фрагментов. Возможно, этот текст уже был в памяти.")
        else:
            await status_msg.edit_text(f"✅ Запомнил текст. Добавлено фрагментов: {inserted_count}")

    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка RAG-памяти: {str(e)}")


async def rag_stats_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    total_chunks, sources = get_rag_stats(chat_id)

    if total_chunks == 0:
        await effective_msg.reply_text("📚 RAG-память пока пустая. Отправь PDF/TXT/DOCX или используй /remember текст.")
        return

    lines = [
        f"📚 В памяти: {total_chunks} фрагментов",
        f"📄 Документов/источников: {len(sources)}",
        "",
        "Источники:",
    ]

    for source_name, count in sources[:15]:
        lines.append(f"• {source_name}: {count} фрагм.")

    if len(sources) > 15:
        lines.append(f"• ...и ещё {len(sources) - 15} источников")

    lines.append("")
    lines.append("Команды:")
    lines.append("• /rag_docs — список документов")
    lines.append("• /rag_delete название_файла — удалить конкретный документ")
    lines.append("• /rag_clear — очистить всю RAG-память этого чата")
    lines.append("• /rag_only_on — отвечать только по документам")
    lines.append("• /rag_only_off — обычный режим")
    lines.append("• /rag_search запрос — поиск по памяти")

    await effective_msg.reply_text("\n".join(lines))


async def rag_docs_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    documents = get_rag_documents(chat_id)

    if not documents:
        await effective_msg.reply_text("📚 В RAG-памяти пока нет документов.")
        return

    lines = ["📚 Документы в памяти:", ""]

    for index, (source_name, chunks_count, first_added_at, last_added_at) in enumerate(documents, start=1):
        last_added_text = time.strftime("%d.%m.%Y %H:%M", time.localtime(last_added_at))
        lines.append(f"{index}. {source_name}")
        lines.append(f"   Фрагментов: {chunks_count} | добавлен/обновлён: {last_added_text}")

    lines.append("")
    lines.append("Удалить документ можно так:")
    lines.append("/rag_delete название_файла")
    lines.append("или по номеру из списка:")
    lines.append("/rag_delete 1")

    await effective_msg.reply_text("\n".join(lines))


async def rag_delete_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    text = effective_msg.text or ""
    parts = text.split(maxsplit=1)

    if len(parts) < 2 or not parts[1].strip():
        await effective_msg.reply_text(
            "Напиши так:\n"
            "/rag_delete название_файла\n\n"
            "Пример:\n"
            "/rag_delete rag_test.pdf\n\n"
            "Можно удалить и по номеру из /rag_docs:\n"
            "/rag_delete 1"
        )
        return

    target = parts[1].strip()
    documents = get_rag_documents(chat_id)

    if not documents:
        await effective_msg.reply_text("📚 В RAG-памяти пока нет документов.")
        return

    # Можно удалить по номеру из /rag_docs
    if target.isdigit():
        doc_index = int(target) - 1

        if doc_index < 0 or doc_index >= len(documents):
            await effective_msg.reply_text(
                f"⚠️ Документа с номером {target} нет. Посмотри список через /rag_docs."
            )
            return

        target = documents[doc_index][0]

    deleted_count, matched_name = delete_rag_document(chat_id, target)

    if deleted_count == 0 or matched_name is None:
        available = "\n".join(f"• {doc_name}" for doc_name, *_ in documents[:10])
        await effective_msg.reply_text(
            "⚠️ Не нашёл такой документ в памяти.\n\n"
            "Проверь название через /rag_docs.\n\n"
            f"Сейчас есть:\n{available}"
        )
        return

    await effective_msg.reply_text(
        f"🗑 Удалил документ из RAG-памяти: {matched_name}\n"
        f"Удалено фрагментов: {deleted_count}"
    )



async def rag_clear_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    deleted_count = clear_rag_memory(chat_id)
    await effective_msg.reply_text(f"🗑 RAG-память очищена. Удалено фрагментов: {deleted_count}")



async def rag_only_on_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    set_rag_only_mode(effective_msg.chat_id, True)
    await effective_msg.reply_text(
        "✅ RAG-only включён.\n"
        "Теперь я буду отвечать только по загруженным документам. Если ответа в памяти нет — так и скажу."
    )


async def rag_only_off_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    set_rag_only_mode(effective_msg.chat_id, False)
    await effective_msg.reply_text(
        "✅ RAG-only выключен.\n"
        "Теперь я снова могу отвечать обычно, а документы использовать как дополнительный контекст."
    )


async def rag_only_status_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    enabled = get_rag_only_mode(effective_msg.chat_id)
    await effective_msg.reply_text(
        "📚 RAG-only: включён" if enabled else "📚 RAG-only: выключен"
    )


async def rag_search_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    text = effective_msg.text or ""
    parts = text.split(maxsplit=1)

    if len(parts) < 2 or not parts[1].strip():
        await effective_msg.reply_text(
            "Напиши так:\n"
            "/rag_search что искать\n\n"
            "Пример:\n"
            "/rag_search код склада"
        )
        return

    query = parts[1].strip()

    try:
        results = rank_rag_chunks(chat_id, query, top_k=RAG_SEARCH_RESULTS, include_below_min=False)
    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка поиска по RAG: {str(e)}")
        return

    if not results:
        await effective_msg.reply_text("📚 Ничего релевантного в RAG-памяти не нашёл.")
        return

    lines = [f"🔎 Поиск по RAG: {query}", ""]
    for index, item in enumerate(results, start=1):
        preview = re.sub(r"\s+", " ", item["text"]).strip()
        if len(preview) > 450:
            preview = preview[:450].rstrip() + "..."

        lines.append(
            f"{index}. {item['source_name']}, фрагмент {item['chunk_index'] + 1}\n"
            f"   Общая: {item['score']:.2f} | semantic: {item['semantic_score']:.2f} | keyword: {item['keyword_score']:.2f}\n"
            f"   {preview}"
        )
        lines.append("")

    await safe_reply_long(effective_msg, "\n".join(lines).strip())


async def admin_stats_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return

    stats = get_admin_stats()
    loaded_models = get_loaded_models()
    lines = [
        "🛠 Админ-статистика",
        "",
        f"Пользователей/чатов в логе: {stats['users_count']}",
        f"Чатов с RAG-памятью: {stats['chats_with_chunks']}",
        f"Документов/источников всего: {stats['total_sources']}",
        f"Фрагментов всего: {stats['total_chunks']}",
        f"RAG-only включён у чатов: {stats['rag_only_count']}",
        f"Размер базы: {stats['db_size']}",
        f"Текущая модель в коде: {currently_loaded_model or 'не определена'}",
        f"Модели, которые LM Studio считает загруженными: {', '.join(loaded_models) if loaded_models else 'нет/не удалось получить'}",
        "",
        f"База: {stats['db_path']}",
    ]
    await safe_reply_long(effective_msg, "\n".join(lines))


async def admin_users_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return

    rows = get_admin_users(limit=20)
    if not rows:
        await effective_msg.reply_text("Пока нет записанных пользователей.")
        return

    lines = ["👥 Последние пользователи/чаты:", ""]
    for index, (chat_id, user_id, username, first_name, last_name, last_seen_at, message_count) in enumerate(rows, start=1):
        last_seen = time.strftime("%d.%m.%Y %H:%M", time.localtime(last_seen_at))
        name = " ".join(part for part in [first_name, last_name] if part).strip() or "без имени"
        username_text = f"@{username}" if username else "без username"
        lines.append(
            f"{index}. {name} ({username_text})\n"
            f"   user_id: {user_id} | chat_id: {chat_id}\n"
            f"   сообщений: {message_count} | был: {last_seen}"
        )
    await safe_reply_long(effective_msg, "\n\n".join(lines))


async def admin_db_size_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return

    await effective_msg.reply_text(
        f"💾 Размер базы: {get_db_size_text()}\n"
        f"Путь: {RAG_DB_PATH}"
    )


async def admin_reload_models_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return

    status_msg = await effective_msg.reply_text("🔄 Синхронизирую загруженные модели LM Studio...")
    sync_loaded_models_on_startup()
    loaded_models = get_loaded_models()
    await status_msg.edit_text(
        "✅ Готово.\n"
        f"Текущая модель: {currently_loaded_model or 'не определена'}\n"
        f"Загружены в LM Studio: {', '.join(loaded_models) if loaded_models else 'нет/не удалось получить'}"
    )


async def handle_document(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg or not effective_msg.document:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    document = effective_msg.document
    file_name = document.file_name or "document"
    extension = os.path.splitext(file_name.lower())[1]

    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        await effective_msg.reply_text(
            "⚠️ Я могу запоминать только PDF, TXT, MD, LOG, CSV, JSON и DOCX."
        )
        return

    if document.file_size and document.file_size > RAG_MAX_FILE_MB * 1024 * 1024:
        await effective_msg.reply_text(
            f"⚠️ Файл слишком большой. Сейчас лимит: {RAG_MAX_FILE_MB} МБ."
        )
        return

    temp_path = None

    try:
        status_msg = await effective_msg.reply_text(f"📄 Читаю файл: {file_name}")

        telegram_file = await document.get_file()
        file_bytes = await telegram_file.download_as_bytearray()

        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name

        await status_msg.edit_text("🧠 Извлекаю текст и добавляю в RAG-память...")

        extracted_text = extract_text_from_file(temp_path, file_name)
        inserted_count = add_text_to_rag(chat_id, file_name, extracted_text)

        if inserted_count == 0:
            await status_msg.edit_text(
                "⚠️ Не нашёл нового текста для запоминания. Возможно, файл пустой, сканированный или уже был добавлен."
            )
            return

        total_chunks, _ = get_rag_stats(chat_id)
        await status_msg.edit_text(
            f"✅ Запомнил документ: {file_name}\n"
            f"Добавлено фрагментов: {inserted_count}\n"
            f"Всего в памяти этого чата: {total_chunks}\n\n"
            "Теперь можешь спрашивать по этому документу обычным сообщением."
        )

        # Если пользователь отправил файл с подписью-вопросом — после запоминания сразу отвечаем на подпись.
        if effective_msg.caption and effective_msg.caption.strip():
            await send_to_model(update, context, effective_msg.caption.strip())

    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка при обработке документа: {str(e)}")

    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


async def handle_message(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id

    if chat_id not in user_selected_model:
        pending_user_content[chat_id] = effective_msg.text
        await ask_model_choice(update, context, reason="restart")
        return

    await send_to_model(update, context, effective_msg.text)


async def _handle_photo_impl(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg or not effective_msg.photo:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id

    try:
        photo_file = await effective_msg.photo[-1].get_file()
        photo_bytes = await photo_file.download_as_bytearray()
        image_b64 = base64.b64encode(photo_bytes).decode("utf-8")

        caption = effective_msg.caption if effective_msg.caption else "Опиши, что на этой картинке"

        user_content = [
            {"type": "text", "text": caption},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}}
        ]

        if chat_id not in user_selected_model:
            pending_user_content[chat_id] = user_content
            await ask_model_choice(update, context, reason="restart")
            return

        await send_to_model(update, context, user_content)

    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка при обработке изображения: {str(e)}")


async def handle_photo(update: Update, context):
    await run_queued(update, context, "обработка изображения", lambda: _handle_photo_impl(update, context))


async def handle_voice(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg or not effective_msg.voice:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id

    try:
        await effective_msg.reply_text("🎙 Распознаю голосовое сообщение...")

        voice_file = await effective_msg.voice.get_file()
        voice_bytes = await voice_file.download_as_bytearray()

        temp_path = f"voice_{chat_id}_{int(time.time())}.ogg"
        with open(temp_path, "wb") as f:
            f.write(voice_bytes)

        segments, info = whisper_model.transcribe(
            temp_path,
            language="ru",
            task="transcribe",
            beam_size=5,
            best_of=5,
            temperature=0.0,
            vad_filter=True,
            condition_on_previous_text=False,
            initial_prompt="Это голосовое сообщение на русском языке. Распознавай речь точно, без перевода."
        )
        recognized_text = " ".join(segment.text for segment in segments).strip()

        os.remove(temp_path)

        if not recognized_text:
            await effective_msg.reply_text("🤷 Не удалось распознать речь, попробуйте ещё раз.")
            return

        await effective_msg.reply_text(f"📝 Распознано: {recognized_text}")

        if chat_id not in user_selected_model:
            pending_user_content[chat_id] = recognized_text
            await ask_model_choice(update, context, reason="restart")
            return

        await send_to_model(update, context, recognized_text)

    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка при распознавании речи: {str(e)}")


async def start(update: Update, context):
    if not update.effective_message:
        return
    if not await check_bot_access(update):
        return
    await update.effective_message.reply_text(
        "👋 Привет! Я локальный AI-ассистент.\n\n"
        "💬 Просто напиши мне что-нибудь, и я отвечу.\n"
        "🖼 Можешь отправить картинку — я её опишу.\n"
        "🎙 Можешь отправить голосовое — я распознаю через Whisper.\n"
        "📚 Можешь отправить PDF/TXT/DOCX — я добавлю документ в RAG-память.\n\n"
        "Основные команды:\n"
        "🔄 /reset — очистить историю диалога\n"
        "🎛 /model — выбрать модель\n"
        "🧠 /remember текст — добавить текст в RAG-память\n"
        "📊 /rag_stats — статистика RAG-памяти\n"
        "📚 /rag_docs — список документов в памяти\n"
        "🗑 /rag_delete название_файла — удалить конкретный документ\n"
        "🧹 /rag_clear — очистить всю RAG-память\n"
        "📚 /rag_only_on — отвечать только по документам\n"
        "🔓 /rag_only_off — обычный режим\n"
        "📚 /rag_only_status — статус RAG-only\n"
        "🔎 /rag_search запрос — поиск по памяти\n"
        "ℹ️ /help — список возможностей"
    )


async def help_command(update: Update, context):
    if not update.effective_message:
        return
    if not await check_bot_access(update):
        return
    await update.effective_message.reply_text(
        "ℹ️ Что я умею:\n\n"
        "💬 Отвечать на вопросы и поддерживать диалог\n"
        "🖼 Анализировать изображения\n"
        "🎙 Распознавать голосовые сообщения через Whisper\n"
        "📚 Запоминать PDF/TXT/DOCX документы через RAG\n"
        "📌 Показывать, какие фрагменты памяти использовал в ответе\n"
        "🔎 Искать по памяти через гибридный поиск: embeddings + keyword\n"
        "📚 Работать в режиме RAG-only, где ответы только по документам\n"
        "🎛 Переключаться между локальными моделями\n\n"
        "Команды:\n"
        "🔄 /reset — очистить историю диалога\n"
        "🎛 /model — выбрать модель\n"
        "🧠 /remember текст — добавить текст в RAG-память\n"
        "📊 /rag_stats — статистика RAG-памяти\n"
        "📚 /rag_docs — список документов в памяти\n"
        "🗑 /rag_delete название_файла — удалить конкретный документ\n"
        "🧹 /rag_clear — очистить всю RAG-память\n"
        "📚 /rag_only_on — включить ответы только по документам\n"
        "🔓 /rag_only_off — выключить RAG-only\n"
        "📚 /rag_only_status — проверить режим\n"
        "🔎 /rag_search запрос — показать найденные фрагменты\n\n"
        "Админ-команды:\n"
        "🛠 /admin_stats — статистика бота\n"
        "👥 /admin_users — последние пользователи\n"
        "💾 /admin_db_size — размер и путь базы\n"
        "🔄 /admin_reload_models — синхронизация LM Studio"
    )


async def reset_history(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    chat_id = effective_msg.chat_id
    conversation_history[chat_id] = [{"role": "system", "content": SYSTEM_PROMPT}]
    await effective_msg.reply_text("🔄 История диалога очищена! Начинаем с чистого листа.")


async def model_command(update: Update, context):
    """Ручной вызов выбора модели через команду /model."""
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return

    chat_id = effective_msg.chat_id
    current_key = user_selected_model.get(chat_id, DEFAULT_MODEL_KEY)

    await effective_msg.reply_text("🎛 Выберите модель:", reply_markup=build_model_keyboard(current_key))


async def model_callback(update: Update, context):
    query = update.callback_query
    await query.answer()
    remember_user_in_db(update)

    if ADMIN_ONLY_MODE and not is_admin(update):
        await query.edit_message_text("⛔ Бот сейчас в закрытом режиме.")
        return

    chat_id = query.message.chat_id
    model_key = query.data.split(":")[1]

    if model_key not in AVAILABLE_MODELS:
        await query.edit_message_text("❌ Неизвестная модель.")
        return

    user_selected_model[chat_id] = model_key
    label = AVAILABLE_MODELS[model_key]["label"]

    await query.edit_message_text(
        f"⏳ Переключаю на: {label}\nЭто может занять 10-30 секунд..."
    )

    model_id = AVAILABLE_MODELS[model_key]["id"]
    ok, error = ensure_model_loaded(model_id)

    if not ok:
        await query.edit_message_text(f"⚠️ Не удалось переключить модель: {error}")
        return

    await query.edit_message_text(f"✅ Активная модель: {label}")

    if chat_id in pending_user_content:
        content = pending_user_content.pop(chat_id)
        await send_to_model(update, context, content)



# =========================================================
# QoL UPGRADE PACK
# Очередь задач, документы с метаданными, кнопки, саммари,
# compare, study, quiz, voice-confirm, backups, export/import, status.
# =========================================================

RAG_FILES_DIR = project_path(os.environ.get("RAG_FILES_DIR", RAG_STORAGE_DIR))
RAG_BACKUPS_DIR = project_path(os.environ.get("RAG_BACKUPS_DIR", "backups"))
AUTO_SUMMARY_ON_UPLOAD = os.environ.get("AUTO_SUMMARY_ON_UPLOAD", "true").strip().lower() in {"1", "true", "yes", "on"}
AUTO_BACKUP_ON_START = os.environ.get("AUTO_BACKUP_ON_START", "true").strip().lower() in {"1", "true", "yes", "on"}
VOICE_CONFIRM_MODE = os.environ.get("VOICE_CONFIRM_MODE", "true").strip().lower() in {"1", "true", "yes", "on"}
MAX_DOC_TEXT_FOR_MODEL = int(os.environ.get("MAX_DOC_TEXT_FOR_MODEL", "9000"))
MAX_BACKUPS_TO_KEEP = int(os.environ.get("MAX_BACKUPS_TO_KEEP", "10"))

chat_task_locks = {}
chat_task_waiting = {}
pending_voice_actions = {}
pending_voice_edit = {}
pending_rag_import = set()


def ensure_sqlite_column(conn, table, column, definition):
    existing = [row[1] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    if column not in existing:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {definition}")


def init_rag_db():
    """Расширенная SQLite-база: chunks + documents + settings + users."""
    migrate_old_rag_db_if_needed()
    os.makedirs(os.path.dirname(os.path.abspath(RAG_DB_PATH)), exist_ok=True)
    os.makedirs(RAG_FILES_DIR, exist_ok=True)
    os.makedirs(RAG_BACKUPS_DIR, exist_ok=True)

    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rag_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT NOT NULL,
                source_name TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                chunk_hash TEXT NOT NULL,
                text TEXT NOT NULL,
                embedding BLOB NOT NULL,
                created_at INTEGER NOT NULL,
                UNIQUE(chat_id, chunk_hash)
            )
            """
        )
        ensure_sqlite_column(conn, "rag_chunks", "document_id", "document_id INTEGER")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_rag_chat_id ON rag_chunks(chat_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_rag_document_id ON rag_chunks(document_id)")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rag_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT NOT NULL,
                source_name TEXT NOT NULL,
                source_hash TEXT NOT NULL,
                saved_path TEXT,
                file_size INTEGER DEFAULT 0,
                file_ext TEXT DEFAULT '',
                chunks_count INTEGER DEFAULT 0,
                summary TEXT,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL,
                UNIQUE(chat_id, source_hash)
            )
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_rag_documents_chat ON rag_documents(chat_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_rag_documents_hash ON rag_documents(chat_id, source_hash)")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS rag_settings (
                chat_id TEXT PRIMARY KEY,
                rag_only INTEGER NOT NULL DEFAULT 0,
                updated_at INTEGER NOT NULL
            )
            """
        )
        ensure_sqlite_column(conn, "rag_settings", "study_mode", "study_mode INTEGER NOT NULL DEFAULT 0")

        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS bot_users (
                chat_id TEXT PRIMARY KEY,
                user_id TEXT,
                username TEXT,
                first_name TEXT,
                last_name TEXT,
                first_seen_at INTEGER NOT NULL,
                last_seen_at INTEGER NOT NULL,
                message_count INTEGER NOT NULL DEFAULT 0
            )
            """
        )
        conn.commit()

    sync_documents_from_old_chunks()


def sync_documents_from_old_chunks():
    """Создаёт записи rag_documents для старых chunk-only документов."""
    with sqlite3.connect(RAG_DB_PATH) as conn:
        rows = conn.execute(
            """
            SELECT chat_id, source_name, COUNT(*), MIN(created_at), MAX(created_at)
            FROM rag_chunks
            WHERE document_id IS NULL
            GROUP BY chat_id, source_name
            """
        ).fetchall()

        for chat_id, source_name, chunks_count, first_at, last_at in rows:
            source_hash = hashlib.sha256(f"legacy:{chat_id}:{source_name}".encode("utf-8")).hexdigest()
            conn.execute(
                """
                INSERT OR IGNORE INTO rag_documents
                (chat_id, source_name, source_hash, saved_path, file_size, file_ext, chunks_count, summary, created_at, updated_at)
                VALUES (?, ?, ?, NULL, 0, ?, ?, NULL, ?, ?)
                """,
                (str(chat_id), source_name, source_hash, os.path.splitext(source_name.lower())[1], chunks_count, first_at or int(time.time()), last_at or int(time.time()))
            )
            doc_id = conn.execute(
                "SELECT id FROM rag_documents WHERE chat_id = ? AND source_hash = ?",
                (str(chat_id), source_hash)
            ).fetchone()[0]
            conn.execute(
                "UPDATE rag_chunks SET document_id = ? WHERE chat_id = ? AND source_name = ? AND document_id IS NULL",
                (doc_id, str(chat_id), source_name)
            )
        conn.commit()


def sanitize_filename(name):
    name = os.path.basename(name or "document")
    name = re.sub(r"[^a-zA-Zа-яА-ЯёЁ0-9._()\- ]+", "_", name).strip()
    return name or "document"


def format_file_size(size):
    try:
        size = int(size or 0)
    except Exception:
        size = 0
    if size < 1024:
        return f"{size} Б"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} КБ"
    return f"{size / (1024 * 1024):.2f} МБ"


def save_original_document(chat_id, file_name, file_bytes, file_hash):
    safe_name = sanitize_filename(file_name)
    chat_dir = os.path.join(RAG_FILES_DIR, f"chat_{chat_id}")
    os.makedirs(chat_dir, exist_ok=True)
    base, ext = os.path.splitext(safe_name)
    saved_name = f"{int(time.time())}_{file_hash[:10]}_{base[:70]}{ext}"
    saved_path = os.path.join(chat_dir, saved_name)
    with open(saved_path, "wb") as f:
        f.write(file_bytes)
    return saved_path


def get_document_by_hash(chat_id, file_hash):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        row = conn.execute(
            """
            SELECT id, source_name, file_size, file_ext, chunks_count, summary, created_at, updated_at, saved_path, source_hash
            FROM rag_documents
            WHERE chat_id = ? AND source_hash = ?
            """,
            (str(chat_id), file_hash)
        ).fetchone()
    return document_row_to_dict(row) if row else None


def document_row_to_dict(row):
    if not row:
        return None
    return {
        "id": row[0],
        "source_name": row[1],
        "file_size": row[2] or 0,
        "file_ext": row[3] or "",
        "chunks_count": row[4] or 0,
        "summary": row[5] or "",
        "created_at": row[6] or 0,
        "updated_at": row[7] or 0,
        "saved_path": row[8] or "",
        "source_hash": row[9] or "",
    }


def get_document_by_id(chat_id, doc_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        row = conn.execute(
            """
            SELECT id, source_name, file_size, file_ext, chunks_count, summary, created_at, updated_at, saved_path, source_hash
            FROM rag_documents
            WHERE chat_id = ? AND id = ?
            """,
            (str(chat_id), int(doc_id))
        ).fetchone()
    return document_row_to_dict(row) if row else None


def get_rag_documents(chat_id):
    """Новый формат: список dict с метаданными документов."""
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        rows = conn.execute(
            """
            SELECT id, source_name, file_size, file_ext, chunks_count, summary, created_at, updated_at, saved_path, source_hash
            FROM rag_documents
            WHERE chat_id = ?
            ORDER BY updated_at DESC, id DESC
            """,
            (str(chat_id),)
        ).fetchall()
    return [document_row_to_dict(row) for row in rows]


def resolve_document(chat_id, target):
    target = (target or "").strip()
    docs = get_rag_documents(chat_id)
    if not target:
        return None
    if target.isdigit():
        idx = int(target) - 1
        if 0 <= idx < len(docs):
            return docs[idx]
        doc = get_document_by_id(chat_id, int(target))
        if doc:
            return doc
    for doc in docs:
        if doc["source_name"] == target:
            return doc
    matches = [doc for doc in docs if doc["source_name"].casefold() == target.casefold()]
    if len(matches) == 1:
        return matches[0]
    partial = [doc for doc in docs if target.casefold() in doc["source_name"].casefold()]
    if len(partial) == 1:
        return partial[0]
    return None


def get_document_chunks(chat_id, doc_id, limit=None):
    init_rag_db()
    sql = """
        SELECT chunk_index, text
        FROM rag_chunks
        WHERE chat_id = ? AND document_id = ?
        ORDER BY chunk_index ASC
    """
    params = [str(chat_id), int(doc_id)]
    if limit:
        sql += " LIMIT ?"
        params.append(int(limit))
    with sqlite3.connect(RAG_DB_PATH) as conn:
        return conn.execute(sql, params).fetchall()


def get_document_text_for_model(chat_id, doc_id, max_chars=MAX_DOC_TEXT_FOR_MODEL):
    chunks = get_document_chunks(chat_id, doc_id)
    parts = []
    total = 0
    for chunk_index, text in chunks:
        block = f"[Фрагмент {chunk_index + 1}]\n{text}"
        if total + len(block) > max_chars:
            remaining = max_chars - total
            if remaining > 300:
                parts.append(block[:remaining].rstrip() + "...")
            break
        parts.append(block)
        total += len(block)
    return "\n\n---\n\n".join(parts)


def add_document_to_rag(chat_id, source_name, text, source_hash=None, saved_path=None, file_size=0, file_ext=None):
    init_rag_db()
    text = clean_text(text)
    chunks = split_text_into_chunks(text)
    if not chunks:
        return None, 0, 0

    source_name = sanitize_filename(source_name) if source_name != "Текст из /remember" else source_name
    source_hash = source_hash or hashlib.sha256(f"{source_name}\n{text}".encode("utf-8")).hexdigest()
    file_ext = file_ext if file_ext is not None else os.path.splitext(source_name.lower())[1]
    now = int(time.time())

    embeddings = embed_texts(chunks)
    inserted_count = 0

    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            """
            INSERT INTO rag_documents
            (chat_id, source_name, source_hash, saved_path, file_size, file_ext, chunks_count, summary, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, 0, NULL, ?, ?)
            ON CONFLICT(chat_id, source_hash) DO UPDATE SET
                source_name = excluded.source_name,
                saved_path = COALESCE(excluded.saved_path, rag_documents.saved_path),
                file_size = CASE WHEN excluded.file_size > 0 THEN excluded.file_size ELSE rag_documents.file_size END,
                file_ext = excluded.file_ext,
                updated_at = excluded.updated_at
            """,
            (str(chat_id), source_name, source_hash, saved_path, int(file_size or 0), file_ext or "", now, now)
        )
        doc_id = conn.execute(
            "SELECT id FROM rag_documents WHERE chat_id = ? AND source_hash = ?",
            (str(chat_id), source_hash)
        ).fetchone()[0]

        for index, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_hash = hashlib.sha256(f"{source_hash}:{index}:{chunk}".encode("utf-8")).hexdigest()
            try:
                conn.execute(
                    """
                    INSERT INTO rag_chunks
                    (chat_id, document_id, source_name, chunk_index, chunk_hash, text, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (str(chat_id), doc_id, source_name, index, chunk_hash, chunk, embedding.astype(np.float32).tobytes(), now)
                )
                inserted_count += 1
            except sqlite3.IntegrityError:
                pass

        chunks_count = conn.execute(
            "SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ? AND document_id = ?",
            (str(chat_id), doc_id)
        ).fetchone()[0]
        conn.execute(
            "UPDATE rag_documents SET chunks_count = ?, updated_at = ? WHERE id = ?",
            (chunks_count, now, doc_id)
        )
        conn.commit()

    return doc_id, inserted_count, len(chunks)


def add_text_to_rag(chat_id, source_name, text):
    source_hash = hashlib.sha256(f"manual:{chat_id}:{source_name}:{text}".encode("utf-8")).hexdigest()
    _, inserted_count, _ = add_document_to_rag(chat_id, source_name, text, source_hash=source_hash, file_ext=".txt")
    return inserted_count


def delete_rag_document(chat_id, target):
    init_rag_db()
    doc = resolve_document(chat_id, str(target))
    if not doc:
        return 0, None
    with sqlite3.connect(RAG_DB_PATH) as conn:
        cursor = conn.execute("DELETE FROM rag_chunks WHERE chat_id = ? AND document_id = ?", (str(chat_id), doc["id"]))
        deleted_count = cursor.rowcount
        conn.execute("DELETE FROM rag_documents WHERE chat_id = ? AND id = ?", (str(chat_id), doc["id"]))
        conn.commit()
    saved_path = doc.get("saved_path") or ""
    if saved_path and os.path.exists(saved_path):
        try:
            os.remove(saved_path)
        except Exception:
            pass
    return deleted_count, doc["source_name"]


def clear_rag_memory(chat_id):
    init_rag_db()
    docs = get_rag_documents(chat_id)
    with sqlite3.connect(RAG_DB_PATH) as conn:
        cursor = conn.execute("DELETE FROM rag_chunks WHERE chat_id = ?", (str(chat_id),))
        deleted_count = cursor.rowcount
        conn.execute("DELETE FROM rag_documents WHERE chat_id = ?", (str(chat_id),))
        conn.commit()
    for doc in docs:
        saved_path = doc.get("saved_path") or ""
        if saved_path and os.path.exists(saved_path):
            try:
                os.remove(saved_path)
            except Exception:
                pass
    return deleted_count


def get_rag_stats(chat_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        total_chunks = conn.execute("SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ?", (str(chat_id),)).fetchone()[0]
        sources = conn.execute(
            """
            SELECT source_name, chunks_count
            FROM rag_documents
            WHERE chat_id = ?
            ORDER BY updated_at DESC
            """,
            (str(chat_id),)
        ).fetchall()
    return total_chunks, sources


def rank_rag_chunks(chat_id, query, top_k=RAG_TOP_K, include_below_min=False):
    if not isinstance(query, str) or len(query.strip()) < 3:
        return []
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        rows = conn.execute(
            """
            SELECT document_id, source_name, chunk_index, text, embedding
            FROM rag_chunks
            WHERE chat_id = ?
            """,
            (str(chat_id),)
        ).fetchall()
    if not rows:
        return []
    query_embedding = embed_texts([query])[0]
    ranked = []
    for document_id, source_name, chunk_index, text, embedding_blob in rows:
        chunk_embedding = np.frombuffer(embedding_blob, dtype=np.float32)
        semantic_score = float(np.dot(query_embedding, chunk_embedding)) if chunk_embedding.size == query_embedding.size else 0.0
        keyword_score = calculate_keyword_score(query, text)
        hybrid_score = (semantic_score * RAG_SEMANTIC_WEIGHT) + (keyword_score * RAG_KEYWORD_WEIGHT)
        if keyword_score >= 0.65:
            hybrid_score = max(hybrid_score, keyword_score)
        if include_below_min or semantic_score >= RAG_MIN_SCORE or keyword_score >= 0.35 or hybrid_score >= RAG_MIN_SCORE:
            ranked.append({
                "score": hybrid_score,
                "semantic_score": semantic_score,
                "keyword_score": keyword_score,
                "source_name": source_name,
                "chunk_index": chunk_index,
                "text": text,
                "document_id": document_id,
            })
    ranked.sort(key=lambda item: item["score"], reverse=True)
    return ranked[:top_k]


def get_chat_setting(chat_id, name, default=False):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        row = conn.execute(f"SELECT {name} FROM rag_settings WHERE chat_id = ?", (str(chat_id),)).fetchone()
    return bool(row and row[0] == 1) if row is not None else bool(default)


def set_chat_setting(chat_id, name, enabled):
    init_rag_db()
    now = int(time.time())
    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            """
            INSERT INTO rag_settings (chat_id, rag_only, study_mode, updated_at)
            VALUES (?, 0, 0, ?)
            ON CONFLICT(chat_id) DO NOTHING
            """,
            (str(chat_id), now)
        )
        conn.execute(f"UPDATE rag_settings SET {name} = ?, updated_at = ? WHERE chat_id = ?", (1 if enabled else 0, now, str(chat_id)))
        conn.commit()


def get_rag_only_mode(chat_id):
    return get_chat_setting(chat_id, "rag_only", False)


def set_rag_only_mode(chat_id, enabled):
    set_chat_setting(chat_id, "rag_only", enabled)


def get_study_mode(chat_id):
    return get_chat_setting(chat_id, "study_mode", False)


def set_study_mode(chat_id, enabled):
    set_chat_setting(chat_id, "study_mode", enabled)


_base_send_to_model = send_to_model
async def send_to_model(update: Update, context, user_content):
    """Обёртка: если включён study_mode, добавляет учебную инструкцию к запросу."""
    try:
        chat_id = update.effective_message.chat_id if update.effective_message else None
        if chat_id and get_study_mode(chat_id) and isinstance(user_content, str):
            user_content = (
                "РЕЖИМ УЧЁБЫ ВКЛЮЧЁН. Отвечай как преподаватель: объясняй понятно, "
                "выделяй главное, задавай 1-2 проверочных вопроса в конце, если это уместно.\n\n"
                f"Запрос пользователя: {user_content}"
            )
    except Exception:
        pass
    await _base_send_to_model(update, context, user_content)


async def run_queued(update: Update, context, task_name, worker):
    """Простая очередь по chat_id: тяжёлые задачи выполняются строго по одной."""
    effective_msg = update.effective_message
    if not effective_msg:
        return
    chat_id = effective_msg.chat_id
    lock = chat_task_locks.setdefault(chat_id, asyncio.Lock())
    queued_msg = None

    if lock.locked():
        chat_task_waiting[chat_id] = chat_task_waiting.get(chat_id, 0) + 1
        position = chat_task_waiting[chat_id]
        queued_msg = await effective_msg.reply_text(f"⏳ Задача поставлена в очередь: {task_name}. Перед вами: {position}")

    async with lock:
        if queued_msg:
            chat_task_waiting[chat_id] = max(0, chat_task_waiting.get(chat_id, 1) - 1)
            try:
                await queued_msg.edit_text(f"▶️ Выполняю задачу: {task_name}")
            except Exception:
                pass
        await worker()


def build_doc_actions_keyboard(doc_id, include_compare=True):
    rows = [
        [InlineKeyboardButton("📌 Резюме", callback_data=f"doc:sum:{doc_id}"), InlineKeyboardButton("⭐ Важное", callback_data=f"doc:keys:{doc_id}")],
        [InlineKeyboardButton("🎓 Учёба", callback_data=f"doc:study:{doc_id}"), InlineKeyboardButton("🧪 Квиз", callback_data=f"doc:quiz:{doc_id}")],
    ]
    if include_compare:
        rows.append([InlineKeyboardButton("⚖️ Сравнить", callback_data=f"doc:cmp:{doc_id}"), InlineKeyboardButton("📎 Получить файл", callback_data=f"doc:get:{doc_id}")])
    else:
        rows.append([InlineKeyboardButton("📎 Получить файл", callback_data=f"doc:get:{doc_id}")])
    rows.append([InlineKeyboardButton("🗑 Удалить", callback_data=f"doc:del:{doc_id}")])
    return InlineKeyboardMarkup(rows)


def format_doc_line(index, doc):
    added = time.strftime("%d.%m.%Y %H:%M", time.localtime(doc["updated_at"] or doc["created_at"] or int(time.time())))
    summary_mark = " | есть резюме" if doc.get("summary") else ""
    file_mark = " | файл сохранён" if doc.get("saved_path") and os.path.exists(doc.get("saved_path")) else ""
    return (
        f"{index}. {doc['source_name']}\n"
        f"   ID: {doc['id']} | {doc['chunks_count']} фрагм. | {format_file_size(doc['file_size'])} | {doc['file_ext'] or 'text'}\n"
        f"   Добавлен/обновлён: {added}{summary_mark}{file_mark}"
    )


async def call_lm_studio_once(chat_id, prompt, system=None, model_id=None, timeout=180):
    model_id = model_id or get_user_model_id(chat_id)
    ok, error = ensure_model_loaded(model_id)
    if not ok:
        raise RuntimeError(f"Проблема с моделью: {error}")
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    else:
        messages.append({"role": "system", "content": SYSTEM_PROMPT})
    messages.append({"role": "user", "content": prompt})
    data = {"model": model_id, "messages": messages, "stream": False}
    response = requests.post(LM_STUDIO_CHAT_URL, headers=get_headers(), json=data, timeout=timeout)
    if response.status_code != 200:
        raise RuntimeError(f"LM Studio API error: {response.status_code} {response.text[:300]}")
    payload = response.json()
    result = payload.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    if PLAIN_TEXT_OUTPUT:
        result = clean_model_output(result)
    return result


def save_document_summary(chat_id, doc_id, summary):
    with sqlite3.connect(RAG_DB_PATH) as conn:
        conn.execute(
            "UPDATE rag_documents SET summary = ?, updated_at = ? WHERE chat_id = ? AND id = ?",
            (summary, int(time.time()), str(chat_id), int(doc_id))
        )
        conn.commit()


async def generate_document_summary(chat_id, doc_id, mode="summary"):
    doc = get_document_by_id(chat_id, doc_id)
    if not doc:
        raise RuntimeError("Документ не найден")
    doc_text = get_document_text_for_model(chat_id, doc_id)
    if not doc_text:
        raise RuntimeError("У документа нет текста в RAG-памяти")

    if mode == "summary":
        task = (
            "Сделай краткое, но полезное резюме документа на русском языке.\n"
            "Структура:\n1) О чём документ\n2) Главные пункты\n3) Важные даты/суммы/имена, если есть\n4) Что стоит запомнить\n"
        )
    elif mode == "keys":
        task = (
            "Вытащи из документа самые важные пункты на русском. "
            "Пиши списком, без воды. Отдельно выдели даты, суммы, имена, риски и обязанности, если они есть."
        )
    elif mode == "study":
        task = (
            "Сделай учебный разбор документа: объясни тему простыми словами, "
            "выдели термины, сделай мини-шпаргалку и 5 вопросов для самопроверки."
        )
    elif mode == "quiz":
        task = (
            "Создай тест по документу на русском: 10 вопросов. "
            "Для каждого вопроса дай варианты A/B/C/D и правильный ответ после всех вопросов."
        )
    else:
        task = "Проанализируй документ на русском языке."

    prompt = f"Документ: {doc['source_name']}\n\n{task}\n\nТЕКСТ ДОКУМЕНТА:\n{doc_text}"
    result = await call_lm_studio_once(chat_id, prompt, system="Ты аккуратно анализируешь документы. Не выдумывай факты вне текста. Пиши обычным текстом без Markdown-разметки.")
    if mode == "summary":
        save_document_summary(chat_id, doc_id, result)
    return result


async def compare_documents(chat_id, doc_id_1, doc_id_2):
    doc1 = get_document_by_id(chat_id, doc_id_1)
    doc2 = get_document_by_id(chat_id, doc_id_2)
    if not doc1 or not doc2:
        raise RuntimeError("Один из документов не найден")
    text1 = get_document_text_for_model(chat_id, doc_id_1, max_chars=5500)
    text2 = get_document_text_for_model(chat_id, doc_id_2, max_chars=5500)
    prompt = (
        f"Сравни два документа на русском языке.\n\n"
        f"Документ 1: {doc1['source_name']}\n{text1}\n\n"
        f"Документ 2: {doc2['source_name']}\n{text2}\n\n"
        "Дай структуру:\n"
        "1) Кратко о каждом документе\n"
        "2) Что совпадает\n"
        "3) Что отличается\n"
        "4) Важные изменения/риски/противоречия\n"
        "5) Итог простыми словами"
    )
    return await call_lm_studio_once(chat_id, prompt, system="Ты сравниваешь документы строго по их тексту. Не выдумывай факты. Пиши обычным текстом без Markdown-разметки.", timeout=240)


def create_rag_backup(reason="manual"):
    init_rag_db()
    os.makedirs(RAG_BACKUPS_DIR, exist_ok=True)
    stamp = time.strftime("%Y%m%d_%H%M%S")
    safe_reason = re.sub(r"[^a-zA-Z0-9_\-]+", "_", reason)[:30]
    backup_path = os.path.join(RAG_BACKUPS_DIR, f"rag_memory_{stamp}_{safe_reason}.sqlite3")
    if os.path.exists(RAG_DB_PATH):
        shutil.copy2(RAG_DB_PATH, backup_path)
    # Чистим старые бэкапы
    backups = sorted(glob.glob(os.path.join(RAG_BACKUPS_DIR, "rag_memory_*.sqlite3")), key=os.path.getmtime, reverse=True)
    for old in backups[MAX_BACKUPS_TO_KEEP:]:
        try:
            os.remove(old)
        except Exception:
            pass
    return backup_path


def list_backups():
    os.makedirs(RAG_BACKUPS_DIR, exist_ok=True)
    return sorted(glob.glob(os.path.join(RAG_BACKUPS_DIR, "rag_memory_*.sqlite3")), key=os.path.getmtime, reverse=True)


def export_rag_memory(chat_id):
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        docs_rows = conn.execute(
            """
            SELECT id, source_name, source_hash, file_size, file_ext, chunks_count, summary, created_at, updated_at
            FROM rag_documents WHERE chat_id = ? ORDER BY updated_at DESC
            """,
            (str(chat_id),)
        ).fetchall()
        chunks_rows = conn.execute(
            """
            SELECT document_id, source_name, chunk_index, chunk_hash, text, embedding, created_at
            FROM rag_chunks WHERE chat_id = ? ORDER BY document_id, chunk_index
            """,
            (str(chat_id),)
        ).fetchall()
    data = {
        "format": "synctech_rag_export_v1",
        "exported_at": int(time.time()),
        "chat_id": str(chat_id),
        "embedding_model": RAG_EMBEDDING_MODEL_NAME,
        "documents": [
            {
                "old_id": row[0], "source_name": row[1], "source_hash": row[2], "file_size": row[3],
                "file_ext": row[4], "chunks_count": row[5], "summary": row[6], "created_at": row[7], "updated_at": row[8]
            } for row in docs_rows
        ],
        "chunks": [
            {
                "old_document_id": row[0], "source_name": row[1], "chunk_index": row[2], "chunk_hash": row[3],
                "text": row[4], "embedding_b64": base64.b64encode(row[5]).decode("ascii"), "created_at": row[6]
            } for row in chunks_rows
        ]
    }
    path = os.path.join(tempfile.gettempdir(), f"rag_export_{chat_id}_{int(time.time())}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return path


def import_rag_memory(chat_id, json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if data.get("format") != "synctech_rag_export_v1":
        raise RuntimeError("Это не поддерживаемый RAG export JSON")
    now = int(time.time())
    old_to_new = {}
    inserted_docs = 0
    inserted_chunks = 0
    init_rag_db()
    with sqlite3.connect(RAG_DB_PATH) as conn:
        for doc in data.get("documents", []):
            source_hash = doc.get("source_hash") or hashlib.sha256(f"import:{doc.get('source_name')}".encode("utf-8")).hexdigest()
            conn.execute(
                """
                INSERT OR IGNORE INTO rag_documents
                (chat_id, source_name, source_hash, saved_path, file_size, file_ext, chunks_count, summary, created_at, updated_at)
                VALUES (?, ?, ?, NULL, ?, ?, 0, ?, ?, ?)
                """,
                (str(chat_id), doc.get("source_name", "imported"), source_hash, int(doc.get("file_size") or 0), doc.get("file_ext") or "", doc.get("summary") or "", doc.get("created_at") or now, now)
            )
            row = conn.execute("SELECT id FROM rag_documents WHERE chat_id = ? AND source_hash = ?", (str(chat_id), source_hash)).fetchone()
            if row:
                old_to_new[doc.get("old_id")] = row[0]
                inserted_docs += 1
        for chunk in data.get("chunks", []):
            new_doc_id = old_to_new.get(chunk.get("old_document_id"))
            emb = base64.b64decode(chunk.get("embedding_b64", ""))
            chunk_hash = hashlib.sha256(f"import:{chat_id}:{chunk.get('chunk_hash')}".encode("utf-8")).hexdigest()
            try:
                conn.execute(
                    """
                    INSERT INTO rag_chunks
                    (chat_id, document_id, source_name, chunk_index, chunk_hash, text, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (str(chat_id), new_doc_id, chunk.get("source_name", "imported"), int(chunk.get("chunk_index") or 0), chunk_hash, chunk.get("text", ""), emb, chunk.get("created_at") or now)
                )
                inserted_chunks += 1
            except sqlite3.IntegrityError:
                pass
        for doc_id in set(old_to_new.values()):
            cnt = conn.execute("SELECT COUNT(*) FROM rag_chunks WHERE chat_id = ? AND document_id = ?", (str(chat_id), doc_id)).fetchone()[0]
            conn.execute("UPDATE rag_documents SET chunks_count = ?, updated_at = ? WHERE id = ?", (cnt, now, doc_id))
        conn.commit()
    return inserted_docs, inserted_chunks


async def advanced_callback(update: Update, context):
    query = update.callback_query
    await query.answer()
    remember_user_in_db(update)
    if ADMIN_ONLY_MODE and not is_admin(update):
        await query.edit_message_text("⛔ Бот сейчас в закрытом режиме.")
        return
    data = query.data or ""
    chat_id = query.message.chat_id

    try:
        if data.startswith("doc:"):
            parts = data.split(":")
            action = parts[1]

            if action in {"sum", "keys", "study", "quiz"}:
                doc_id = int(parts[2])
                mode_map = {"sum": "summary", "keys": "keys", "study": "study", "quiz": "quiz"}
                title_map = {"sum": "📌 Резюме", "keys": "⭐ Важные пункты", "study": "🎓 Учебный разбор", "quiz": "🧪 Квиз"}
                await query.edit_message_text(f"{title_map[action]}: готовлю...")
                result = await generate_document_summary(chat_id, doc_id, mode=mode_map[action])
                await safe_reply_long(query.message, f"{title_map[action]}\n\n{result}")
                doc = get_document_by_id(chat_id, doc_id)
                if doc:
                    await query.message.reply_text("Действия с документом:", reply_markup=build_doc_actions_keyboard(doc_id))
                return

            if action == "cmp":
                doc_id = int(parts[2])
                docs = [doc for doc in get_rag_documents(chat_id) if doc["id"] != doc_id]
                if not docs:
                    await query.edit_message_text("⚠️ Для сравнения нужен хотя бы ещё один документ в памяти.")
                    return
                rows = []
                for doc in docs[:10]:
                    rows.append([InlineKeyboardButton(doc["source_name"][:45], callback_data=f"doc:cmp2:{doc_id}:{doc['id']}")])
                await query.edit_message_text("⚖️ Выбери второй документ для сравнения:", reply_markup=InlineKeyboardMarkup(rows))
                return

            if action == "cmp2":
                doc1 = int(parts[2]); doc2 = int(parts[3])
                await query.edit_message_text("⚖️ Сравниваю документы...")
                result = await compare_documents(chat_id, doc1, doc2)
                await safe_reply_long(query.message, f"⚖️ Сравнение документов\n\n{result}")
                return

            if action == "get":
                doc = get_document_by_id(chat_id, int(parts[2]))
                if not doc or not doc.get("saved_path") or not os.path.exists(doc["saved_path"]):
                    await query.edit_message_text("⚠️ Оригинальный файл не найден. Возможно, это был текст из /remember или старый документ без сохранённого файла.")
                    return
                await query.message.reply_document(document=open(doc["saved_path"], "rb"), filename=doc["source_name"])
                return

            if action == "del":
                doc_id = int(parts[2])
                doc = get_document_by_id(chat_id, doc_id)
                if not doc:
                    await query.edit_message_text("⚠️ Документ уже не найден.")
                    return
                deleted, name = delete_rag_document_by_id(chat_id, doc_id)
                await query.edit_message_text(f"🗑 Удалил документ: {name}\nУдалено фрагментов: {deleted}")
                return

        if data.startswith("voice:"):
            parts = data.split(":")
            action = parts[1]
            token = parts[2]
            text = pending_voice_actions.get(token)
            if not text:
                await query.edit_message_text("⚠️ Этот текст уже устарел. Отправь голосовое заново.")
                return
            if action == "send":
                await query.edit_message_text(f"📝 Отправляю модели:\n{text}")
                await send_to_model(update, context, text)
                return
            if action == "short":
                await query.edit_message_text(f"📝 Делаю краткий ответ по тексту:\n{text}")
                await send_to_model(update, context, f"Ответь кратко на это голосовое сообщение:\n{text}")
                return
            if action == "remember":
                inserted = add_text_to_rag(chat_id, "Голосовая заметка", text)
                await query.edit_message_text(f"🧠 Запомнил голосовую заметку. Добавлено фрагментов: {inserted}")
                return
            if action == "edit":
                pending_voice_edit[chat_id] = token
                await query.edit_message_text("✏️ Ок, отправь следующим сообщением исправленный текст. Я покажу кнопки ещё раз.")
                return
            if action == "cancel":
                pending_voice_actions.pop(token, None)
                await query.edit_message_text("❌ Голосовое отменено.")
                return
    except Exception as e:
        await query.message.reply_text(f"❌ Ошибка действия: {str(e)}")


async def _handle_document_impl(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg or not effective_msg.document:
        return
    if not await check_bot_access(update):
        return
    chat_id = effective_msg.chat_id
    document = effective_msg.document
    file_name = document.file_name or "document"
    extension = os.path.splitext(file_name.lower())[1]

    # Импорт RAG export JSON через /rag_import
    if chat_id in pending_rag_import:
        pending_rag_import.discard(chat_id)
        if extension != ".json":
            await effective_msg.reply_text("⚠️ Для импорта нужен JSON-файл, созданный командой /rag_export.")
            return
        temp_path = None
        try:
            status = await effective_msg.reply_text("📥 Импортирую RAG-память...")
            telegram_file = await document.get_file()
            file_bytes = await telegram_file.download_as_bytearray()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            docs_count, chunks_count = import_rag_memory(chat_id, temp_path)
            await status.edit_text(f"✅ Импорт завершён. Документов: {docs_count}, фрагментов: {chunks_count}")
        except Exception as e:
            await effective_msg.reply_text(f"❌ Ошибка импорта: {str(e)}")
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)
        return

    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        await effective_msg.reply_text("⚠️ Я могу запоминать только PDF, TXT, MD, LOG, CSV, JSON и DOCX.")
        return
    if document.file_size and document.file_size > RAG_MAX_FILE_MB * 1024 * 1024:
        await effective_msg.reply_text(f"⚠️ Файл слишком большой. Сейчас лимит: {RAG_MAX_FILE_MB} МБ.")
        return

    temp_path = None
    try:
        status_msg = await effective_msg.reply_text(f"📄 Читаю файл: {file_name}")
        telegram_file = await document.get_file()
        file_bytes = await telegram_file.download_as_bytearray()
        file_hash = hashlib.sha256(file_bytes).hexdigest()

        duplicate = get_document_by_hash(chat_id, file_hash)
        if duplicate:
            await status_msg.edit_text(
                f"⚠️ Такой документ уже был загружен: {duplicate['source_name']}\n"
                f"ID: {duplicate['id']} | фрагментов: {duplicate['chunks_count']}\n\n"
                "Чтобы загрузить заново, сначала удали старый документ."
            )
            await effective_msg.reply_text("Действия с уже загруженным документом:", reply_markup=build_doc_actions_keyboard(duplicate["id"]))
            return

        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name

        await status_msg.edit_text("🧠 Извлекаю текст и добавляю в RAG-память...")
        extracted_text = extract_text_from_file(temp_path, file_name)
        saved_path = save_original_document(chat_id, file_name, file_bytes, file_hash)
        doc_id, inserted_count, total_chunks = add_document_to_rag(
            chat_id,
            file_name,
            extracted_text,
            source_hash=file_hash,
            saved_path=saved_path,
            file_size=len(file_bytes),
            file_ext=extension,
        )

        if inserted_count == 0 or not doc_id:
            await status_msg.edit_text("⚠️ Не нашёл нового текста для запоминания. Возможно, файл пустой, сканированный или уже был добавлен.")
            return

        total_in_chat, _ = get_rag_stats(chat_id)
        await status_msg.edit_text(
            f"✅ Запомнил документ: {file_name}\n"
            f"ID: {doc_id}\n"
            f"Добавлено фрагментов: {inserted_count}\n"
            f"Всего в памяти этого чата: {total_in_chat}\n\n"
            "Готовлю кнопки действий..."
        )

        if AUTO_SUMMARY_ON_UPLOAD:
            try:
                summary_msg = await effective_msg.reply_text("📌 Делаю краткое резюме документа...")
                summary = await generate_document_summary(chat_id, doc_id, mode="summary")
                await safe_edit(summary_msg, f"📌 Авто-резюме: {file_name}\n\n{summary}")
            except Exception as e:
                await effective_msg.reply_text(f"⚠️ Документ запомнен, но авто-резюме не получилось: {str(e)}")

        await effective_msg.reply_text("Что сделать с документом?", reply_markup=build_doc_actions_keyboard(doc_id))

        if effective_msg.caption and effective_msg.caption.strip():
            await send_to_model(update, context, effective_msg.caption.strip())

    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка при обработке документа: {str(e)}")
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


async def handle_document(update: Update, context):
    await run_queued(update, context, "обработка документа", lambda: _handle_document_impl(update, context))


async def _handle_message_impl(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    chat_id = effective_msg.chat_id

    if chat_id in pending_voice_edit:
        token = pending_voice_edit.pop(chat_id)
        corrected = (effective_msg.text or "").strip()
        if not corrected:
            await effective_msg.reply_text("⚠️ Пустой текст. Голосовое действие отменено.")
            return
        pending_voice_actions[token] = corrected
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("✅ Ответить", callback_data=f"voice:send:{token}"), InlineKeyboardButton("✏️ Исправить", callback_data=f"voice:edit:{token}")],
            [InlineKeyboardButton("🧠 Запомнить", callback_data=f"voice:remember:{token}"), InlineKeyboardButton("⚡ Кратко", callback_data=f"voice:short:{token}")],
            [InlineKeyboardButton("❌ Отмена", callback_data=f"voice:cancel:{token}")],
        ])
        await effective_msg.reply_text(f"📝 Исправленный текст:\n{corrected}", reply_markup=keyboard)
        return

    if chat_id not in user_selected_model:
        pending_user_content[chat_id] = effective_msg.text
        await ask_model_choice(update, context, reason="restart")
        return
    await send_to_model(update, context, effective_msg.text)


async def handle_message(update: Update, context):
    await run_queued(update, context, "текстовый запрос", lambda: _handle_message_impl(update, context))


async def _handle_voice_impl(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg or not effective_msg.voice:
        return
    if not await check_bot_access(update):
        return
    chat_id = effective_msg.chat_id
    temp_path = None
    try:
        status = await effective_msg.reply_text("🎙 Распознаю голосовое сообщение...")
        voice_file = await effective_msg.voice.get_file()
        voice_bytes = await voice_file.download_as_bytearray()
        temp_path = os.path.join(tempfile.gettempdir(), f"voice_{chat_id}_{int(time.time())}.ogg")
        with open(temp_path, "wb") as f:
            f.write(voice_bytes)
        segments, info = whisper_model.transcribe(
            temp_path,
            language="ru",
            task="transcribe",
            beam_size=5,
            best_of=5,
            temperature=0.0,
            vad_filter=True,
            condition_on_previous_text=False,
            initial_prompt="Это голосовое сообщение на русском языке. Распознавай речь точно, без перевода."
        )
        recognized_text = " ".join(segment.text for segment in segments).strip()
        if not recognized_text:
            await status.edit_text("🤷 Не удалось распознать речь, попробуйте ещё раз.")
            return

        token = uuid.uuid4().hex[:12]
        pending_voice_actions[token] = recognized_text
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("✅ Ответить", callback_data=f"voice:send:{token}"), InlineKeyboardButton("✏️ Исправить", callback_data=f"voice:edit:{token}")],
            [InlineKeyboardButton("🧠 Запомнить", callback_data=f"voice:remember:{token}"), InlineKeyboardButton("⚡ Кратко", callback_data=f"voice:short:{token}")],
            [InlineKeyboardButton("❌ Отмена", callback_data=f"voice:cancel:{token}")],
        ])
        await status.edit_text(f"📝 Распознано:\n{recognized_text}")
        if VOICE_CONFIRM_MODE:
            await effective_msg.reply_text("Что сделать с распознанным текстом?", reply_markup=keyboard)
        else:
            await send_to_model(update, context, recognized_text)
    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка при распознавании речи: {str(e)}")
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


async def handle_voice(update: Update, context):
    await run_queued(update, context, "распознавание голосового", lambda: _handle_voice_impl(update, context))


async def rag_docs_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    docs = get_rag_documents(effective_msg.chat_id)
    if not docs:
        await effective_msg.reply_text("📚 В RAG-памяти пока нет документов.")
        return
    lines = ["📚 Документы в памяти:", ""]
    for i, doc in enumerate(docs, start=1):
        lines.append(format_doc_line(i, doc))
    lines.append("")
    lines.append("Команды: /rag_summary 1, /quiz 1, /rag_compare 1 2, /rag_delete 1")
    await safe_reply_long(effective_msg, "\n\n".join(lines))
    # Кнопки для последних 5 документов, чтобы не засорять чат.
    for doc in docs[:5]:
        await effective_msg.reply_text(f"Действия: {doc['source_name']}", reply_markup=build_doc_actions_keyboard(doc["id"]))


async def rag_delete_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    parts = (effective_msg.text or "").split(maxsplit=1)
    if len(parts) < 2:
        await effective_msg.reply_text("Напиши так:\n/rag_delete название_файла\nили:\n/rag_delete 1")
        return
    deleted, name = delete_rag_document(effective_msg.chat_id, parts[1].strip())
    if deleted == 0 or not name:
        await effective_msg.reply_text("⚠️ Не нашёл такой документ. Посмотри список через /rag_docs.")
        return
    await effective_msg.reply_text(f"🗑 Удалил документ: {name}\nУдалено фрагментов: {deleted}")


async def rag_summary_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    parts = (effective_msg.text or "").split(maxsplit=1)
    if len(parts) < 2:
        await effective_msg.reply_text("Напиши так:\n/rag_summary 1\nили:\n/rag_summary название_файла")
        return
    doc = resolve_document(effective_msg.chat_id, parts[1].strip())
    if not doc:
        await effective_msg.reply_text("⚠️ Документ не найден. Посмотри /rag_docs.")
        return
    if doc.get("summary"):
        await safe_reply_long(effective_msg, f"📌 Резюме: {doc['source_name']}\n\n{doc['summary']}")
        return
    msg = await effective_msg.reply_text("📌 Делаю резюме...")
    try:
        summary = await generate_document_summary(effective_msg.chat_id, doc["id"], mode="summary")
        await safe_edit(msg, f"📌 Резюме: {doc['source_name']}\n\n{summary}")
    except Exception as e:
        await safe_edit(msg, f"❌ Ошибка резюме: {str(e)}")


async def rag_compare_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    args = (effective_msg.text or "").split(maxsplit=2)
    if len(args) < 3:
        await effective_msg.reply_text("Напиши так:\n/rag_compare 1 2\nили:\n/rag_compare file1.pdf file2.pdf")
        return
    doc1 = resolve_document(effective_msg.chat_id, args[1])
    doc2 = resolve_document(effective_msg.chat_id, args[2])
    if not doc1 or not doc2:
        await effective_msg.reply_text("⚠️ Один из документов не найден. Посмотри /rag_docs.")
        return
    msg = await effective_msg.reply_text("⚖️ Сравниваю документы...")
    try:
        result = await compare_documents(effective_msg.chat_id, doc1["id"], doc2["id"])
        await safe_edit(msg, f"⚖️ Сравнение: {doc1['source_name']} ↔ {doc2['source_name']}\n\n{result}")
    except Exception as e:
        await safe_edit(msg, f"❌ Ошибка сравнения: {str(e)}")


async def study_mode_on_command(update: Update, context):
    if update.effective_message and await check_bot_access(update):
        set_study_mode(update.effective_message.chat_id, True)
        await update.effective_message.reply_text("🎓 Study mode включён. Теперь я буду отвечать как преподаватель и добавлять проверочные вопросы.")


async def study_mode_off_command(update: Update, context):
    if update.effective_message and await check_bot_access(update):
        set_study_mode(update.effective_message.chat_id, False)
        await update.effective_message.reply_text("🎓 Study mode выключен.")


async def study_mode_status_command(update: Update, context):
    if update.effective_message and await check_bot_access(update):
        enabled = get_study_mode(update.effective_message.chat_id)
        await update.effective_message.reply_text("🎓 Study mode: включён" if enabled else "🎓 Study mode: выключен")


async def quiz_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    parts = (effective_msg.text or "").split(maxsplit=1)
    if len(parts) >= 2:
        doc = resolve_document(effective_msg.chat_id, parts[1].strip())
    else:
        docs = get_rag_documents(effective_msg.chat_id)
        doc = docs[0] if docs else None
    if not doc:
        await effective_msg.reply_text("⚠️ Документ не найден. Используй /rag_docs и потом /quiz 1")
        return
    msg = await effective_msg.reply_text("🧪 Генерирую квиз...")
    try:
        result = await generate_document_summary(effective_msg.chat_id, doc["id"], mode="quiz")
        await safe_edit(msg, f"🧪 Квиз: {doc['source_name']}\n\n{result}")
    except Exception as e:
        await safe_edit(msg, f"❌ Ошибка квиза: {str(e)}")


async def backup_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return
    try:
        path = create_rag_backup("manual")
        await effective_msg.reply_document(document=open(path, "rb"), filename=os.path.basename(path), caption=f"✅ Бэкап создан: {os.path.basename(path)}")
    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка бэкапа: {str(e)}")


async def backups_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await require_admin(update):
        return
    backups = list_backups()
    if not backups:
        await effective_msg.reply_text("💾 Бэкапов пока нет.")
        return
    lines = ["💾 Последние бэкапы:", ""]
    for i, path in enumerate(backups[:10], start=1):
        lines.append(f"{i}. {os.path.basename(path)} — {format_file_size(os.path.getsize(path))}")
    await effective_msg.reply_text("\n".join(lines))


async def rag_export_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    try:
        path = export_rag_memory(effective_msg.chat_id)
        await effective_msg.reply_document(document=open(path, "rb"), filename=os.path.basename(path), caption="📤 Экспорт RAG-памяти этого чата")
    except Exception as e:
        await effective_msg.reply_text(f"❌ Ошибка экспорта: {str(e)}")


async def rag_import_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    pending_rag_import.add(effective_msg.chat_id)
    await effective_msg.reply_text("📥 Отправь следующим сообщением JSON-файл, созданный командой /rag_export.")


async def status_command(update: Update, context):
    effective_msg = update.effective_message
    if not effective_msg:
        return
    if not await check_bot_access(update):
        return
    chat_id = effective_msg.chat_id
    lines = ["🩺 Status / healthcheck", ""]
    lines.append("✅ Telegram bot: работает")
    try:
        loaded = get_loaded_models()
        lines.append(f"✅ LM Studio: доступен | загружено моделей: {len(loaded)}")
    except Exception as e:
        lines.append(f"❌ LM Studio: ошибка {e}")
    try:
        _ = whisper_model
        lines.append(f"✅ Whisper: {WHISPER_MODEL_NAME}")
    except Exception as e:
        lines.append(f"❌ Whisper: {e}")
    try:
        init_rag_db()
        total_chunks, docs = get_rag_stats(chat_id)
        lines.append(f"✅ RAG DB: доступна | {get_db_size_text()}")
        lines.append(f"📚 В этом чате: документов {len(docs)}, фрагментов {total_chunks}")
        lines.append(f"📍 DB: {RAG_DB_PATH}")
        lines.append(f"📁 Files: {RAG_FILES_DIR}")
    except Exception as e:
        lines.append(f"❌ RAG DB: {e}")
    if SentenceTransformer is None:
        lines.append("❌ sentence-transformers: не установлен")
    else:
        lines.append("✅ sentence-transformers: установлен")
    lines.append("🎓 Study mode: включён" if get_study_mode(chat_id) else "🎓 Study mode: выключен")
    lines.append("📚 RAG-only: включён" if get_rag_only_mode(chat_id) else "📚 RAG-only: выключен")
    await effective_msg.reply_text("\n".join(lines))


async def start(update: Update, context):
    if not update.effective_message:
        return
    if not await check_bot_access(update):
        return
    await update.effective_message.reply_text(
        "👋 Привет! Я локальный AI-ассистент.\n\n"
        "💬 Текст, 🖼 картинки, 🎙 голосовые, 📚 PDF/TXT/DOCX RAG-память.\n\n"
        "Главное:\n"
        "/model — выбрать модель\n"
        "/status — healthcheck\n"
        "/rag_docs — документы с кнопками\n"
        "/rag_summary 1 — резюме документа\n"
        "/rag_compare 1 2 — сравнить документы\n"
        "/quiz 1 — квиз по документу\n"
        "/study_mode_on — режим учёбы\n"
        "/rag_export — экспорт памяти\n"
        "/rag_import — импорт памяти\n"
        "/backup — бэкап базы, только админ\n"
        "/help — все команды"
    )


async def help_command(update: Update, context):
    if not update.effective_message:
        return
    if not await check_bot_access(update):
        return
    await update.effective_message.reply_text(
        "ℹ️ Возможности:\n\n"
        "✅ Очередь задач по чату\n"
        "✅ Кнопки после загрузки документа\n"
        "✅ Авто-резюме документа\n"
        "✅ Сравнение документов\n"
        "✅ Study mode\n"
        "✅ Квиз по документу\n"
        "✅ Голосовое: ответить / исправить / запомнить / кратко\n"
        "✅ /rag_docs с датами, размером и кнопками\n"
        "✅ Сохранение оригинальных файлов\n"
        "✅ Бэкапы, экспорт/импорт памяти\n"
        "✅ Анти-дубли по hash файла\n"
        "✅ /status healthcheck\n\n"
        "Команды:\n"
        "/remember текст\n"
        "/rag_stats\n"
        "/rag_docs\n"
        "/rag_search запрос\n"
        "/rag_summary 1\n"
        "/rag_compare 1 2\n"
        "/rag_delete 1\n"
        "/rag_clear\n"
        "/rag_only_on | /rag_only_off | /rag_only_status\n"
        "/study_mode_on | /study_mode_off | /study_mode_status\n"
        "/quiz 1\n"
        "/rag_export\n"
        "/rag_import\n"
        "/status\n\n"
        "Админ:\n"
        "/admin_stats\n/admin_users\n/admin_db_size\n/admin_reload_models\n/backup\n/backups"
    )

if __name__ == "__main__":
    if not TELEGRAM_BOT_TOKEN:
        print("Ошибка: TELEGRAM_BOT_TOKEN не установлен")
        exit(1)

    print("Проверяю состояние LM Studio при старте...")
    sync_loaded_models_on_startup()

    print("Инициализирую RAG-память...")
    print(f"RAG база: {RAG_DB_PATH}")
    if ADMIN_TELEGRAM_IDS:
        print(f"Админы: {', '.join(sorted(ADMIN_TELEGRAM_IDS))}")
    if ADMIN_ONLY_MODE:
        print("Включён ADMIN_ONLY_MODE: бот доступен только админам")
    init_rag_db()
    if AUTO_BACKUP_ON_START:
        try:
            backup_path = create_rag_backup("auto_start")
            print(f"Авто-бэкап RAG создан: {backup_path}")
        except Exception as e:
            print(f"Не удалось создать авто-бэкап RAG: {e}")

    application = ApplicationBuilder().token(TELEGRAM_BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("reset", reset_history))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("remember", remember_command))
    application.add_handler(CommandHandler("rag_stats", rag_stats_command))
    application.add_handler(CommandHandler("rag_docs", rag_docs_command))
    application.add_handler(CommandHandler("rag_delete", rag_delete_command))
    application.add_handler(CommandHandler("rag_clear", rag_clear_command))
    application.add_handler(CommandHandler("rag_only_on", rag_only_on_command))
    application.add_handler(CommandHandler("rag_only_off", rag_only_off_command))
    application.add_handler(CommandHandler("rag_only_status", rag_only_status_command))
    application.add_handler(CommandHandler("rag_search", rag_search_command))
    application.add_handler(CommandHandler("rag_summary", rag_summary_command))
    application.add_handler(CommandHandler("rag_compare", rag_compare_command))
    application.add_handler(CommandHandler("quiz", quiz_command))
    application.add_handler(CommandHandler("study_mode_on", study_mode_on_command))
    application.add_handler(CommandHandler("study_mode_off", study_mode_off_command))
    application.add_handler(CommandHandler("study_mode_status", study_mode_status_command))
    application.add_handler(CommandHandler("rag_export", rag_export_command))
    application.add_handler(CommandHandler("rag_import", rag_import_command))
    application.add_handler(CommandHandler("status", status_command))
    application.add_handler(CommandHandler("admin_stats", admin_stats_command))
    application.add_handler(CommandHandler("admin_users", admin_users_command))
    application.add_handler(CommandHandler("admin_db_size", admin_db_size_command))
    application.add_handler(CommandHandler("admin_reload_models", admin_reload_models_command))
    application.add_handler(CommandHandler("backup", backup_command))
    application.add_handler(CommandHandler("backups", backups_command))
    application.add_handler(CommandHandler("model", model_command))
    application.add_handler(CallbackQueryHandler(advanced_callback, pattern="^(doc|voice):"))
    application.add_handler(CallbackQueryHandler(model_callback, pattern="^model:"))
    application.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(MessageHandler(filters.VOICE, handle_voice))

    print("Бот запущен. Ожидание сообщений...")
    application.run_polling()